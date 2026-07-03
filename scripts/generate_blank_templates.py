# -*- coding: utf-8 -*-
import os
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

# Margins: Top: 2cm, Bottom: 2cm, Left: 3cm, Right: 2cm
def set_margins(doc):
    s = doc.sections[0]
    s.top_margin    = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin   = Cm(3.0)
    s.right_margin  = Cm(2.0)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    new_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        new_borders.append(border)
    tblPr.append(new_borders)

FONT_NAME = "Times New Roman"

def add_run(para, text, bold=False, italic=False, size=12, underline=False):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    r.font.name = FONT_NAME
    r.font.size = Pt(size)
    return r

def add_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False, size=12, space_before=0, space_after=6, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size)
    return p

def add_national_header(doc, company_name_placeholder, doc_number_placeholder):
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_table)
    header_table.columns[0].width = Cm(7.0)
    header_table.columns[1].width = Cm(9.0)
    
    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]
    
    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.paragraph_format.space_after = Pt(2)
    add_run(p_left, company_name_placeholder, bold=True, size=10.0)
    
    p_no = left_cell.add_paragraph()
    p_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_no.paragraph_format.space_after = Pt(0)
    add_run(p_no, f"Số: {doc_number_placeholder}", size=10.5)
    
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.paragraph_format.space_after = Pt(2)
    add_run(p_right, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, size=11)
    
    p_motto = right_cell.add_paragraph()
    p_motto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_motto.paragraph_format.space_after = Pt(4)
    add_run(p_motto, "Độc lập - Tự do - Hạnh phúc", bold=True, size=11.5)
    
    p_line = right_cell.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_after = Pt(0)
    add_run(p_line, "_______________", bold=True, size=11)

def add_signature_block(doc, rep_role_placeholder, show_reporter=False):
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(sig_table)
    sig_table.columns[0].width = Cm(7.5)
    sig_table.columns[1].width = Cm(8.5)
    
    s_left = sig_table.rows[0].cells[0]
    s_right = sig_table.rows[0].cells[1]
    
    if show_reporter:
        p_rep_title = s_left.paragraphs[0]
        p_rep_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rep_title.paragraph_format.space_after = Pt(2)
        add_run(p_rep_title, "NGƯỜI ĐÁNH GIÁ", bold=True)
        
        p_rep_sign = s_left.add_paragraph()
        p_rep_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rep_sign.paragraph_format.space_after = Pt(40)
        add_run(p_rep_sign, "(Ký và ghi rõ họ tên)", italic=True, size=11)
        
        p_rep_name = s_left.add_paragraph()
        p_rep_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p_rep_name, "[Họ và tên người đánh giá]", bold=True)

    p_date = s_right.paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_after = Pt(2)
    add_run(p_date, "[Địa danh], ngày .... tháng .... năm 20...", italic=True)
    
    p_rep = s_right.add_paragraph()
    p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rep.paragraph_format.space_after = Pt(2)
    add_run(p_rep, "ĐẠI DIỆN TỔ CHỨC, CÁ NHÂN" if not show_reporter else "XÁC NHẬN CỦA LÃNH ĐẠO", bold=True)
    
    p_role = s_right.add_paragraph()
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role.paragraph_format.space_after = Pt(40)
    add_run(p_role, f"(Ký tên, {rep_role_placeholder}, đóng dấu)", italic=True, size=11)
    
    p_name_sign = s_right.add_paragraph()
    p_name_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_name_sign, "[Họ và tên người đại diện]", bold=True)

# ---------------------------------------------------------------------------
#  1. PHU LUC IV: BAN DANG KY CONG BO HOP QUY
# ---------------------------------------------------------------------------
def build_template_phuluc_iv(filepath):
    doc = Document()
    set_margins(doc)
    
    add_national_header(doc, "[TÊN DOANH NGHIỆP/CÁ NHÂN]", "..../ĐKCBHQ")
    add_paragraph(doc, "", space_after=18)
    add_paragraph(doc, "BẢN ĐĂNG KÝ CÔNG BỐ HỢP QUY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=12)

    p_info = add_paragraph(doc, "Tên tổ chức, cá nhân: ", bold=True, space_after=4)
    add_run(p_info, "[Điền tên đầy đủ của doanh nghiệp/cá nhân nhập khẩu]")
    
    p_addr = add_paragraph(doc, "Địa chỉ: ", bold=True, space_after=4)
    add_run(p_addr, "[Điền địa chỉ trụ sở chính đăng ký đăng ký kinh doanh]")
    
    p_contact = add_paragraph(doc, "Điện thoại: ", bold=True, space_after=12)
    add_run(p_contact, "[Số điện thoại]      ")
    add_run(p_contact, "Fax: ", bold=True)
    add_run(p_contact, "[Số Fax]      ")
    add_run(p_contact, "E-mail: ", bold=True)
    add_run(p_contact, "[Địa chỉ Email]")

    add_paragraph(doc, "CÔNG BỐ:", bold=True, size=12.5, space_after=6)
    add_paragraph(doc, "Sản phẩm, hàng hóa, quá trình, dịch vụ, môi trường hoặc các đối tượng khác trong hoạt động kinh tế - xã hội:", italic=True, space_after=6)
    
    p_name = add_paragraph(doc, "- Tên gọi sản phẩm: ", bold=True, space_after=4)
    add_run(p_name, "[Tên gọi sản phẩm thương mại hoặc kỹ thuật]")
    
    p_model = add_paragraph(doc, "- Kiểu, loại/Model: ", bold=True, space_after=4)
    add_run(p_model, "Chi tiết tại Danh mục sản phẩm kèm theo (Phụ lục số 01)", italic=True)
    
    p_brand = add_paragraph(doc, "- Nhãn hiệu/Hiệu: ", bold=True, space_after=4)
    add_run(p_brand, "Chi tiết tại Danh mục sản phẩm kèm theo (Phụ lục số 01)", italic=True)
    
    p_specs = add_paragraph(doc, "- Đặc trưng kỹ thuật: ", bold=True, space_after=8)
    add_run(p_specs, "Chi tiết tại Danh mục sản phẩm kèm theo (Phụ lục số 01)", italic=True)

    p_conform = add_paragraph(doc, "Phù hợp với quy chuẩn kỹ thuật:", bold=True, space_after=4)
    p_conform_detail = add_paragraph(doc, "- Số hiệu, ký hiệu, tên gọi: ", bold=True, space_after=8)
    add_run(p_conform_detail, "[Điền số hiệu và tên Quy chuẩn kỹ thuật áp dụng, ví dụ: QCVN 04:2009/BKHCN]")

    add_paragraph(doc, "Thông tin bổ sung:", bold=True, space_after=4)
    p_basis = add_paragraph(doc, "- Căn cứ công bố hợp quy: ", bold=True, space_after=4)
    add_run(p_basis, "[Điền căn cứ, ví dụ: Giấy chứng nhận hợp quy số ... hoặc Kết quả tự đánh giá hợp quy của doanh nghiệp]")
    
    p_method = add_paragraph(doc, "- Phương thức đánh giá sự phù hợp: ", bold=True, space_after=8)
    add_run(p_method, "[Điền phương thức áp dụng, ví dụ: Phương thức 1/Phương thức 5/Phương thức 7 quy định tại Thông tư 14/2026/TT-BKHCN]")

    add_paragraph(doc, "Loại hình đánh giá:", bold=True, space_after=4)
    
    add_paragraph(doc, "   [ ]  Tổ chức chứng nhận đánh giá:", bold=True, space_after=2)
    add_paragraph(doc, "         + Tên tổ chức chứng nhận: [Tên tổ chức chứng nhận đánh giá, ví dụ: QUACERT]", space_after=2)
    add_paragraph(doc, "         + Số giấy chứng nhận: [Số giấy chứng nhận hợp quy]", space_after=2)
    add_paragraph(doc, "         + Ngày cấp: [Ngày cấp giấy chứng nhận]", space_after=6)
    
    p_eval2 = add_paragraph(doc, "   [ ]  Tự đánh giá: ", bold=True, space_after=12)
    add_run(p_eval2, "Ngày lãnh đạo tổ chức, cá nhân ký xác nhận Báo cáo tự đánh giá: ..../..../20...")

    p_commit = add_paragraph(doc, space_after=18)
    p_commit.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p_commit, "[Tên doanh nghiệp/cá nhân] ", bold=True)
    add_run(p_commit, "cam kết và chịu trách nhiệm về tính phù hợp của sản phẩm, hàng hóa, quá trình, dịch vụ, môi trường và các đối tượng khác trong hoạt động kinh tế - xã hội do mình sản xuất, kinh doanh, bảo quản, vận chuyển, sử dụng, khai thác.")
    
    add_paragraph(doc, "", space_after=12)
    add_signature_block(doc, "chức vụ", show_reporter=False)
    doc.save(filepath)
    print(f"Created template: {filepath}")

# ---------------------------------------------------------------------------
#  2. PHU LUC V: BAO CAO TU DANH GIA HOP QUY
# ---------------------------------------------------------------------------
def build_template_phuluc_v(filepath):
    doc = Document()
    set_margins(doc)
    
    add_national_header(doc, "[TÊN DOANH NGHIỆP/CÁ NHÂN]", "..../BC-CBHQ")
    add_paragraph(doc, "", space_after=18)
    add_paragraph(doc, "BÁO CÁO TỰ ĐÁNH GIÁ HỢP QUY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=12)

    add_paragraph(doc, "1. Tên tổ chức, cá nhân; địa chỉ; điện thoại, fax:", bold=True, space_after=3)
    p1_desc = add_paragraph(doc, "   - Tên tổ chức: ", bold=True, space_after=2)
    add_run(p1_desc, "[Điền tên đầy đủ doanh nghiệp]")
    p1_addr = add_paragraph(doc, "   - Địa chỉ: ", bold=True, space_after=2)
    add_run(p1_addr, "[Điền địa chỉ đăng ký kinh doanh]")
    p1_contact = add_paragraph(doc, "   - Điện thoại: ", bold=True, space_after=8)
    add_run(p1_contact, "[Điện thoại]       Fax: [Fax]       Email: [Email]")

    add_paragraph(doc, "2. Tên sản phẩm, hàng hoá:", bold=True, space_after=3)
    p2_desc = add_paragraph(doc, "   - Tên hàng hóa: ", bold=True, space_after=2)
    add_run(p2_desc, "[Tên sản phẩm]")
    p2_models = add_paragraph(doc, "   - Kiểu, loại/Model: ", bold=True, space_after=2)
    add_run(p2_models, "Chi tiết tại Danh mục sản phẩm kèm theo (Phụ lục số 01)")
    p2_brand = add_paragraph(doc, "   - Nhãn hiệu: ", bold=True, space_after=2)
    add_run(p2_brand, "[Nhãn hiệu/Thương hiệu sản phẩm]")
    p2_specs = add_paragraph(doc, "   - Thông số kỹ thuật chính: ", bold=True, space_after=8)
    add_run(p2_specs, "Chi tiết tại Danh mục sản phẩm kèm theo (Phụ lục số 01)", italic=True)

    add_paragraph(doc, "3. Địa điểm, ngày đánh giá:", bold=True, space_after=3)
    p3_desc = add_paragraph(doc, "   - Địa điểm thực hiện đánh giá: ", bold=True, space_after=2)
    add_run(p3_desc, "[Kho lưu trữ bảo quản hoặc nhà máy sản xuất]")
    p3_date = add_paragraph(doc, "   - Ngày đánh giá: ", bold=True, space_after=8)
    add_run(p3_date, ".... tháng .... năm 20...")

    add_paragraph(doc, "4. Số hiệu tiêu chuẩn/quy chuẩn kỹ thuật áp dụng:", bold=True, space_after=3)
    p4_qcvn = add_paragraph(doc, "   - ", space_after=8)
    add_run(p4_qcvn, "[Số hiệu QCVN, ví dụ: QCVN 04:2009/BKHCN]", bold=True)
    add_run(p4_qcvn, " - [Tên quy chuẩn kỹ thuật áp dụng]")

    add_paragraph(doc, "5. Tên tổ chức thử nghiệm sản phẩm, hàng hoá:", bold=True, space_after=3)
    p5_desc = add_paragraph(doc, "   - Tên phòng thử nghiệm: ", bold=True, space_after=2)
    add_run(p5_desc, "[Tên phòng thử nghiệm thực hiện test mẫu chỉ định/công nhận]")
    p5_vilas = add_paragraph(doc, "   - Mã số công nhận/chỉ định: ", bold=True, space_after=8)
    add_run(p5_vilas, "[Số hiệu VILAS và quyết định chỉ định của cơ quan nhà nước]")

    add_paragraph(doc, "6. Đánh giá về kết quả thử nghiệm theo tiêu chuẩn/quy chuẩn kỹ thuật áp dụng:", bold=True, space_after=3)
    p6_desc = add_paragraph(doc, space_after=8)
    p6_desc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    add_run(p6_desc, "   - Đối với Tự đánh giá: Căn cứ Phiếu kết quả thử nghiệm số [Số phiếu test] ngày [Ngày cấp] của [Tên Lab] cấp cho mẫu hàng hóa [Tên sản phẩm].\n", italic=True)
    add_run(p6_desc, "   - Đối với Tổ chức chứng nhận đánh giá: Căn cứ Giấy chứng nhận hợp quy số [Số GCN] ngày [Ngày cấp] do [Tổ chức chứng nhận] cấp (trên cơ sở Phiếu kết quả thử nghiệm số [Số phiếu test] ngày [Ngày cấp] của [Tên Lab]).\n", italic=True)
    add_run(p6_desc, "   Đánh giá chung: Các chỉ tiêu kỹ thuật đều phù hợp giới hạn quy định và đạt chất lượng phù hợp với quy chuẩn áp dụng.")

    add_paragraph(doc, "7. Các nội dung khác (nếu có):", bold=True, space_after=3)
    p7_desc = add_paragraph(doc, space_after=8)
    p7_desc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p7_desc, "   - Biện pháp kiểm soát chất lượng: Duy trì kiểm soát hồ sơ kỹ thuật, hồ sơ nhập khẩu hải quan chính ngạch. Thực hiện dán nhãn hợp quy CR và nhãn phụ tiếng Việt đầy đủ trước khi lưu thông.\n")
    add_run(p7_desc, "   - Kế hoạch giám sát: Lấy mẫu thử nghiệm và kiểm soát chất lượng định kỳ hàng năm.")

    add_paragraph(doc, "8. Kết luận:", bold=True, space_after=3)
    add_paragraph(doc, "   [ ]  Sản phẩm, hàng hoá phù hợp tiêu chuẩn/quy chuẩn kỹ thuật.", bold=True, space_after=2)
    add_paragraph(doc, "   [ ]  Sản phẩm, hàng hoá không phù hợp tiêu chuẩn/quy chuẩn kỹ thuật.", bold=True, space_after=12)

    p_commit = add_paragraph(doc, space_after=18)
    p_commit.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p_commit, "[Tên doanh nghiệp/cá nhân] cam kết chất lượng sản phẩm, hàng hóa phù hợp với quy chuẩn kỹ thuật, tiêu chuẩn công bố áp dụng và hoàn toàn chịu trách nhiệm trước pháp luật về chất lượng sản phẩm, hàng hóa và kết quả tự đánh giá.")

    add_signature_block(doc, "chức vụ", show_reporter=True)
    doc.save(filepath)
    print(f"Created template: {filepath}")

# ---------------------------------------------------------------------------
#  3. DANH MUC SAN PHAM 9 COT (No Weight Column)
# ---------------------------------------------------------------------------
def build_template_danhmuc(filepath):
    doc = Document()
    set_margins(doc)
    
    add_national_header(doc, "[TÊN DOANH NGHIỆP/CÁ NHÂN]", "Kèm theo số: [Số Bản đăng ký]")
    add_paragraph(doc, "", space_after=12)
    
    p_pl = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)
    add_run(p_pl, "PHỤ LỤC SỐ 01", bold=True, size=11)
    
    add_paragraph(doc, "DANH MỤC SẢN PHẨM [TÊN DÒNG HÀNG]", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_after=2)
    
    p_sub = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_run(p_sub, "(Kèm theo bản đăng ký công bố hợp quy số: .................... và báo cáo tự đánh giá số: ....................)", bold=True, italic=True, size=10.5)

    tbl_headers = [
        "TT", "Tên sản phẩm, hàng hóa", "Nhãn hiệu", "Mã HS", "Xuất xứ", 
        "Số lượng (cái)", "Đặc tính kỹ thuật", "Nhóm sản phẩm", "Quy chuẩn áp dụng"
    ]
    tbl_widths = [0.8, 4.2, 1.3, 1.5, 1.3, 1.3, 3.2, 1.2, 2.5]
    
    table = doc.add_table(rows=2, cols=len(tbl_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for col_idx, (text, width) in enumerate(zip(tbl_headers, tbl_widths)):
        cell = table.rows[0].cells[col_idx]
        cell.width = Cm(width)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, text, bold=True, size=9.5)
        
    # Placeholder row
    row_data = [
        "1",
        "[Tên sản phẩm thương mại đầy đủ, kèm model, mới 100%]",
        "[Nhãn hiệu]",
        "[Mã HS]",
        "[Xuất xứ]",
        "[Số lượng]",
        "[Đặc trưng thông số kỹ thuật chính: điện áp, công suất...]",
        "Nhóm 2",
        "[Số hiệu QCVN]"
    ]
    
    for col_idx, text in enumerate(row_data):
        cell = table.rows[1].cells[col_idx]
        cell.width = Cm(tbl_widths[col_idx])
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if col_idx in [0, 2, 3, 4, 5, 7, 8]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, text, size=9.0)
            
    add_paragraph(doc, "", space_after=18)
    
    sig_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(sig_table)
    sig_table.columns[0].width = Cm(8.0)
    sig_table.columns[1].width = Cm(8.0)
    
    s_right = sig_table.rows[0].cells[1]
    p_rep = s_right.paragraphs[0]
    p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_rep, "ĐẠI DIỆN TỔ CHỨC, CÁ NHÂN", bold=True, size=11)
    
    p_role = s_right.add_paragraph()
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role.paragraph_format.space_after = Pt(40)
    add_run(p_role, "(Ký tên, chức vụ, đóng dấu)", italic=True, size=10)
    
    p_name = s_right.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_name, "[Họ và tên người đại diện]", bold=True, size=11)
    
    doc.save(filepath)
    print(f"Created template: {filepath}")

# ---------------------------------------------------------------------------
#  4. MAU 01 PHU LUC VII: DANG KY KIEM TRA NHAT NUOC VE CHAT LUONG (ND 37)
# ---------------------------------------------------------------------------
def build_template_kiemtrachatlương(filepath):
    doc = Document()
    set_margins(doc)
    
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_table)
    header_table.columns[0].width = Cm(7.5)
    header_table.columns[1].width = Cm(8.5)
    
    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]
    
    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_left, "[TÊN DOANH NGHIỆP/CÁ NHÂN ĐĂNG KÝ]", bold=True, size=10.0)
    
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_right, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, size=11)
    
    p_motto = right_cell.add_paragraph()
    p_motto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_motto.paragraph_format.space_after = Pt(4)
    add_run(p_motto, "Độc lập - Tự do - Hạnh phúc", bold=True, size=11.5)
    
    p_line = right_cell.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_line, "_______________", bold=True, size=11)
    
    add_paragraph(doc, "", space_after=18)
    add_paragraph(doc, "BẢN ĐĂNG KÝ KIỂM TRA NHÀ NƯỚC VỀ CHẤT LƯỢNG HÀNG HÓA NHẬP KHẨU", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_after=4)
    add_paragraph(doc, "(Mẫu số 01 Phụ lục VII ban hành kèm theo Nghị định số 37/2026/NĐ-CP)", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10.5, space_after=12)

    p_to = add_paragraph(doc, "Kính gửi: ", bold=True, space_after=12)
    add_run(p_to, "[Điền tên Cơ quan kiểm tra: Ủy ban Tiêu chuẩn Đo lường Chất lượng Quốc gia hoặc Chi cục TCĐLCL địa phương]")

    add_paragraph(doc, "Tổ chức, cá nhân đăng ký:", bold=True, space_after=4)
    add_run(add_paragraph(doc, "1. Tên tổ chức, cá nhân nhập khẩu: ", bold=True, space_after=2), "[Điền tên đầy đủ doanh nghiệp/cá nhân]")
    add_run(add_paragraph(doc, "2. Địa chỉ: ", bold=True, space_after=2), "[Địa chỉ trụ sở chính]")
    
    p_contact = add_paragraph(doc, "3. Điện thoại: ", bold=True, space_after=8)
    add_run(p_contact, "[Điện thoại]      ")
    add_run(p_contact, "Fax: ", bold=True)
    add_run(p_contact, "[Fax]      ")
    add_run(p_contact, "E-mail: ", bold=True)
    add_run(p_contact, "[Email]")

    add_paragraph(doc, "Đăng ký kiểm tra chất lượng đối với lô hàng nhập khẩu sau đây:", bold=True, space_after=4)
    
    # Lot details table
    tbl_headers = [
        "TT", "Tên hàng hóa", "Nhãn hiệu, kiểu loại", "Đặc tính kỹ thuật", "Xuất xứ, Nhà SX", "Mã HS", "Số lượng", "Chứng từ kèm theo"
    ]
    tbl_widths = [0.6, 2.5, 2.0, 2.5, 2.0, 1.2, 1.2, 4.0]
    table = doc.add_table(rows=2, cols=len(tbl_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for col_idx, (text, width) in enumerate(zip(tbl_headers, tbl_widths)):
        cell = table.rows[0].cells[col_idx]
        cell.width = Cm(width)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_run(p, text, bold=True, size=9.0)
        
    placeholder_row = [
        "1",
        "[Tên sản phẩm]",
        "[Nhãn hiệu, model/mã loại]",
        "[Công suất, điện áp...]",
        "[Nước xuất xứ, tên Nhà SX]",
        "[Mã HS]",
        "[Số lượng/ Khối lượng]",
        "Hợp đồng số: ...\nHóa đơn số: ...\nVận đơn số: ...\nTờ khai số: ..."
    ]
    
    for col_idx, text in enumerate(placeholder_row):
        cell = table.rows[1].cells[col_idx]
        cell.width = Cm(tbl_widths[col_idx])
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if col_idx in [0, 5, 6]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, text, size=8.5)

    add_paragraph(doc, "", space_after=12)
    add_run(add_paragraph(doc, "4. Địa điểm bảo quản hàng hóa sau khi thông quan/đưa về kho: ", bold=True, space_after=2), "[Điền địa chỉ kho bảo quản của doanh nghiệp (Ví dụ: Kho Eureka)]")
    add_run(add_paragraph(doc, "5. Thời gian kiểm tra chất lượng dự kiến: ", bold=True, space_after=12), "[Điền khoảng thời gian dự kiến thực hiện đánh giá phù hợp lấy mẫu]")

    add_paragraph(doc, "Tài liệu gửi kèm theo hồ sơ đăng ký:", bold=True, space_after=4)
    add_paragraph(doc, "   [ ]  Bản sao Hợp đồng mua bán (Contract) / Hóa đơn thương mại (Invoice).", space_after=2)
    add_paragraph(doc, "   [ ]  Danh mục hàng hóa (Packing List) / Vận đơn (Bill of Lading).", space_after=2)
    add_paragraph(doc, "   [ ]  Bản sao Giấy chứng nhận hợp quy (nếu đã được cấp chứng nhận kiểu loại dài hạn).", space_after=2)
    add_paragraph(doc, "   [ ]  Ảnh chụp sản phẩm, hình vẽ cấu tạo, mẫu nhãn chính và nhãn phụ dự thảo.", space_after=2)
    add_paragraph(doc, "   [ ]  Tờ khai hải quan hàng hóa nhập khẩu.", space_after=12)

    p_commit = add_paragraph(doc, space_after=18)
    p_commit.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p_commit, "Chúng tôi cam kết chịu trách nhiệm hoàn toàn trước pháp luật về tính hợp pháp, hợp lệ của các chứng từ trong hồ sơ đăng ký kiểm tra chất lượng và cam kết đưa sản phẩm lưu thông đạt chất lượng phù hợp quy chuẩn kỹ thuật.")

    add_signature_block(doc, "chức vụ", show_reporter=False)
    doc.save(filepath)
    print(f"Created template: {filepath}")


def main():
    target_dir = r"d:\AI AGENT THUY\AI agent Kinh doanh\templates\biểu mẫu giấy tờ\.10. MẪU GIẤY TỜ LÀM KIỂM TRA CHUYÊN NGÀNH\MẪU HỒ SƠ LÀM KIỂM TRA CHUYÊN NGÀNH\BỘ HỒ SƠ LÀM KIỂM TRA CHẤT LƯỢNG\BỘ TEMPLATE CHUẨN AGENT 5"
    os.makedirs(target_dir, exist_ok=True)
    
    f1 = os.path.join(target_dir, "PHU_LUC_IV_BAN_DANG_KY_CONG_BO_HOP_QUY.docx")
    f2 = os.path.join(target_dir, "PHU_LUC_V_BAO_CAO_TU_DANH_GIA_HOP_QUY.docx")
    f3 = os.path.join(target_dir, "DANH_MUC_SAN_PHAM_9_COT.docx")
    f4 = os.path.join(target_dir, "MAU_01_PHU_LUC_VII_DANG_KY_KIEM_TRA_CHAT_LUONG.docx")
    
    build_template_phuluc_iv(f1)
    build_template_phuluc_v(f2)
    build_template_danhmuc(f3)
    build_template_kiemtrachatlương(f4)
    
    print("FINISHED creating blank conformity templates for Agent 5.")

if __name__ == "__main__":
    main()
