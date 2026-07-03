# -*- coding: utf-8 -*-
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

# Page Margins: Top: 2cm, Bottom: 2cm, Left: 3cm, Right: 2cm
def set_margins(doc):
    s = doc.sections[0]
    s.top_margin    = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin   = Cm(3.0)
    s.right_margin  = Cm(2.0)

def create_element(name):
    return OxmlElement(name)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    new_borders = create_element('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = create_element(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        new_borders.append(border)
    tblPr.append(new_borders)

FONT_NAME = "Times New Roman"

def add_run(para, text, bold=False, italic=False, size=12, color_rgb=None, underline=False):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    r.font.name = FONT_NAME
    r.font.size = Pt(size)
    if color_rgb:
        r.font.color.rgb = color_rgb
    return r

def add_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False, size=12, space_before=0, space_after=6, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size)
    return p

def add_national_header(doc, doc_number="05/ĐKCBHQ-ERK"):
    # Header block: Importer info & National Motto (using a borderless 2-column table)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_table)
    
    header_table.columns[0].width = Cm(6.5)
    header_table.columns[1].width = Cm(9.5)
    
    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]
    
    # Left Column: Company Name & Document No.
    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.paragraph_format.space_after = Pt(2)
    p_left.paragraph_format.line_spacing = 1.15
    add_run(p_left, "CÔNG TY TNHH THƯƠNG MẠI\nXUẤT NHẬP KHẨU EUREKA", bold=True, size=10.5)
    
    p_no = left_cell.add_paragraph()
    p_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_no.paragraph_format.space_after = Pt(0)
    add_run(p_no, f"Số: {doc_number}", size=11, bold=False)
    
    # Right Column: National Motto
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.paragraph_format.space_after = Pt(2)
    add_run(p_right, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, size=11)
    
    p_motto = right_cell.add_paragraph()
    p_motto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_motto.paragraph_format.space_after = Pt(4)
    add_run(p_motto, "Độc lập - Tự do - Hạnh phúc", bold=True, size=11.5)
    
    # Line under right column motto
    p_line = right_cell.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(0)
    add_run(p_line, "_______________", bold=True, size=11)

def add_signature_block(doc, show_reporter=False):
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(sig_table)
    sig_table.columns[0].width = Cm(7.0)
    sig_table.columns[1].width = Cm(9.0)
    
    s_left = sig_table.rows[0].cells[0]
    s_right = sig_table.rows[0].cells[1]
    
    if show_reporter:
        # Left side: Reporter signature
        p_rep_title = s_left.paragraphs[0]
        p_rep_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rep_title.paragraph_format.space_after = Pt(2)
        add_run(p_rep_title, "NGƯỜI ĐÁNH GIÁ", bold=True)
        
        p_rep_sign = s_left.add_paragraph()
        p_rep_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rep_sign.paragraph_format.space_after = Pt(40)
        add_run(p_rep_sign, "(Ký và ghi rõ họ tên)", italic=True, size=11)
        
        p_rep_name = s_left.add_paragraph()
        p_rep_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p_rep_name, "Nguyễn Văn Hùng", bold=True)

    # Right side: Representative signature
    p_date = s_right.paragraphs[0] if not show_reporter else s_right.paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_after = Pt(2)
    add_run(p_date, "Hà Nội, ngày 02 tháng 07 năm 2026", italic=True)
    
    p_rep = s_right.add_paragraph()
    p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rep.paragraph_format.space_after = Pt(2)
    add_run(p_rep, "XÁC NHẬN CỦA LÃNH ĐẠO" if show_reporter else "ĐẠI DIỆN TỔ CHỨC, CÁ NHÂN", bold=True)
    
    p_role = s_right.add_paragraph()
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role.paragraph_format.space_after = Pt(40)
    add_run(p_role, "(Ký tên, chức vụ, đóng dấu)", italic=True, size=11)
    
    p_name_sign = s_right.add_paragraph()
    p_name_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name_sign.paragraph_format.space_after = Pt(0)
    add_run(p_name_sign, "Trần Đức Thắng", bold=True)

# ---------------------------------------------------------------------------
#  1. BAN DANG KY CONG BO HOP QUY (Kem Danh muc)
# ---------------------------------------------------------------------------
def build_ban_dang_ky():
    doc = Document()
    set_margins(doc)
    add_national_header(doc, "05/ĐKCBHQ-ERK")
    
    add_paragraph(doc, "", space_after=18)
    add_paragraph(doc, "BẢN ĐĂNG KÝ CÔNG BỐ HỢP QUY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=12)

    p_info_intro = add_paragraph(doc, "Tên tổ chức, cá nhân: ", bold=True, space_after=4)
    add_run(p_info_intro, "CÔNG TY TNHH THƯƠNG MẠI XUẤT NHẬP KHẨU EUREKA", bold=True)
    
    p_addr = add_paragraph(doc, "Địa chỉ: ", bold=True, space_after=4)
    add_run(p_addr, "Số 3, hẻm 56, ngõ An Sơn, phố Đại La, phường Trương Định, quận Hai Bà Trưng, thành phố Hà Nội, Việt Nam")
    
    p_contact = add_paragraph(doc, "Điện thoại: ", bold=True, space_after=12)
    add_run(p_contact, "024.3999.8888      ")
    add_run(p_contact, "Fax: ", bold=True)
    add_run(p_contact, "024.3999.8889      ")
    add_run(p_contact, "E-mail: ", bold=True)
    add_run(p_contact, "contact@eurekalogistics.vn")

    add_paragraph(doc, "ĐĂNG KÝ CÔNG BỐ:", bold=True, size=12.5, space_after=6)
    add_paragraph(doc, "Sản phẩm, hàng hóa, quá trình, dịch vụ, môi trường hoặc các đối tượng khác trong hoạt động kinh tế - xã hội:", italic=True, space_after=6)
    
    p_name = add_paragraph(doc, "- Tên gọi sản phẩm: ", bold=True, space_after=4)
    add_run(p_name, "Máy sấy tóc dùng trong gia dụng")
    
    p_model = add_paragraph(doc, "- Kiểu, loại/Model: ", bold=True, space_after=4)
    add_run(p_model, "Chi tiết tại Danh mục sản phẩm kèm theo Bản đăng ký công bố hợp quy (Phụ lục số 01)", italic=True)
    
    p_brand = add_paragraph(doc, "- Nhãn hiệu/Hiệu: ", bold=True, space_after=4)
    add_run(p_brand, "Chi tiết tại Danh mục sản phẩm kèm theo Bản đăng ký công bố hợp quy (Phụ lục số 01)", italic=True)
    
    p_specs = add_paragraph(doc, "- Đặc trưng kỹ thuật: ", bold=True, space_after=8)
    add_run(p_specs, "Chi tiết tại Danh mục sản phẩm kèm theo Bản đăng ký công bố hợp quy (Phụ lục số 01)", italic=True)

    p_conform = add_paragraph(doc, "Phù hợp với quy chuẩn kỹ thuật:", bold=True, space_after=4)
    p_conform_detail = add_paragraph(doc, "- Số hiệu, ký hiệu, tên gọi: ", bold=True, space_after=8)
    add_run(p_conform_detail, "QCVN 04:2009/BKHCN và Sửa đổi 1:2016 QCVN 04:2009/BKHCN - Quy chuẩn kỹ thuật quốc gia về an toàn đối với thiết bị điện và điện tử; QCVN 09:2012/BKHCN và Sửa đổi 1:2018 QCVN 09:2012/BKHCN - Quy chuẩn kỹ thuật quốc gia về tương thích điện từ đối với thiết bị điện và điện tử gia dụng.")

    add_paragraph(doc, "Thông tin bổ sung:", bold=True, space_after=4)
    p_basis = add_paragraph(doc, "- Căn cứ công bố hợp quy: ", bold=True, space_after=4)
    add_run(p_basis, "Kết quả chứng nhận hợp quy của Tổ chức chứng nhận được chỉ định kết hợp Báo cáo tự đánh giá của doanh nghiệp.")
    
    p_method = add_paragraph(doc, "- Phương thức đánh giá sự phù hợp: ", bold=True, space_after=8)
    add_run(p_method, "Phương thức 5 và Phương thức 1 theo quy định tại Thông tư số 14/2026/TT-BKHCN.")

    add_paragraph(doc, "Loại hình đánh giá:", bold=True, space_after=4)
    
    p_eval1 = add_paragraph(doc, "   [x]  Tổ chức chứng nhận đánh giá (đối với an toàn QCVN 04):", bold=True, space_after=2)
    p_eval1_detail = add_paragraph(doc, "         + Tên tổ chức chứng nhận: ", bold=True, space_after=2)
    add_run(p_eval1_detail, "Trung tâm Chứng nhận Phù hợp (QUACERT)")
    p_eval1_cert = add_paragraph(doc, "         + Số giấy chứng nhận: ", bold=True, space_after=2)
    add_run(p_eval1_cert, "26.0702/GCNHQ-QUACERT")
    p_eval1_date = add_paragraph(doc, "         + Ngày cấp: ", bold=True, space_after=6)
    add_run(p_eval1_date, "30 tháng 06 năm 2026")
    
    p_eval2 = add_paragraph(doc, "   [x]  Tự đánh giá (đối với tương thích điện từ EMC QCVN 09): ", bold=True, space_after=12)
    add_run(p_eval2, "Ngày lãnh đạo tổ chức, cá nhân ký xác nhận Báo cáo tự đánh giá: 02 tháng 07 năm 2026")

    p_commit = add_paragraph(doc, space_after=18)
    p_commit.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p_commit, "CÔNG TY TNHH THƯƠNG MẠI XUẤT NHẬP KHẨU EUREKA ", bold=True)
    add_run(p_commit, "cam kết và chịu trách nhiệm về tính phù hợp của sản phẩm, hàng hóa, quá trình, dịch vụ, môi trường và các đối tượng khác trong hoạt động kinh tế - xã hội do mình sản xuất, kinh doanh, bảo quản, vận chuyển, sử dụng, khai thác.")

    add_signature_block(doc, show_reporter=False)
    return doc

# ---------------------------------------------------------------------------
#  2. DANH MUC SAN PHAM (Phu luc 01)
# ---------------------------------------------------------------------------
def build_danh_muc_san_pham():
    doc = Document()
    set_margins(doc)
    
    # Simple National Header
    add_national_header(doc, "Kèm theo số: 05/ĐKCBHQ-ERK")
    add_paragraph(doc, "", space_after=18)
    
    # Title
    add_paragraph(doc, "DANH MỤC SẢN PHẨM MÁY SẤY TÓC", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_after=2)
    add_paragraph(doc, "(KÈM THEO BẢN ĐĂNG KÝ CÔNG BỐ HỢP QUY SỐ: 05/ĐKCBHQ-ERK)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, italic=True, size=11, space_after=12)

    # Table columns definition
    tbl_headers = [
        "TT", "Tên sản phẩm, hàng hóa", "Nhãn hiệu", "Mã HS", "Xuất xứ", 
        "Số lượng (cái)", "Đặc tính kỹ thuật", "Nhóm SP", "Quy chuẩn áp dụng", "Khối lượng (kg)"
    ]
    tbl_widths = [0.8, 3.2, 1.2, 1.3, 1.1, 1.1, 2.7, 1.1, 2.3, 1.2]
    
    rows_data = [
        [
            "1", "Máy sấy tóc dùng trong gia dụng, Model A1, mới 100%", "A", "8516.31.00", "Trung Quốc",
            "500", "Điện áp: 220V~50Hz;\nCông suất: 1600W", "Nhóm 2", "QCVN 04:2009/BKHCN\nQCVN 09:2012/BKHCN", "250.0"
        ],
        [
            "2", "Máy sấy tóc dùng trong gia dụng, Model A2, mới 100%", "A", "8516.31.00", "Trung Quốc",
            "300", "Điện áp: 220V~50Hz;\nCông suất: 1800W", "Nhóm 2", "QCVN 04:2009/BKHCN\nQCVN 09:2012/BKHCN", "180.0"
        ],
        [
            "3", "Máy sấy tóc dùng trong gia dụng, Model A3, mới 100%", "A", "8516.31.00", "Trung Quốc",
            "200", "Điện áp: 220V~50Hz;\nCông suất: 2000W", "Nhóm 2", "QCVN 04:2009/BKHCN\nQCVN 09:2012/BKHCN", "130.0"
        ]
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=len(tbl_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style Header
    for col_idx, (text, width) in enumerate(zip(tbl_headers, tbl_widths)):
        cell = table.rows[0].cells[col_idx]
        cell.width = Cm(width)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, text, bold=True, size=9.5)
        
    # Style Rows
    for row_idx, row_data in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = Cm(tbl_widths[col_idx])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.0
            
            # Alignments: columns like TT, HS, Origin, Qty, Group, Standard, Weight centered. Name and Specs left-aligned.
            if col_idx in [0, 2, 3, 4, 5, 7, 8, 9]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            add_run(p, text, size=9.0)
            
    add_paragraph(doc, "", space_after=18)
    
    # Signature block: centered on the right
    sig_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(sig_table)
    sig_table.columns[0].width = Cm(8.0)
    sig_table.columns[1].width = Cm(8.0)
    
    s_right = sig_table.rows[0].cells[1]
    
    p_rep = s_right.paragraphs[0]
    p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rep.paragraph_format.space_after = Pt(2)
    add_run(p_rep, "ĐẠI DIỆN TỔ CHỨC, CÁ NHÂN", bold=True, size=11)
    
    p_role = s_right.add_paragraph()
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role.paragraph_format.space_after = Pt(40)
    add_run(p_role, "(Ký tên, chức vụ, đóng dấu)", italic=True, size=10)
    
    p_name = s_right.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_name, "Trần Đức Thắng", bold=True, size=11)
    
    return doc

# ---------------------------------------------------------------------------
#  3. BAO CAO TU DANH GIA HOP QUY (Phu luc V)
# ---------------------------------------------------------------------------
def build_bao_cao_tu_danh_gia():
    doc = Document()
    set_margins(doc)
    
    # Standard Header
    add_national_header(doc, "05/BC-ERK")
    add_paragraph(doc, "", space_after=18)
    
    # Title
    add_paragraph(doc, "BÁO CÁO TỰ ĐÁNH GIÁ HỢP QUY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=12)

    # 1. Company Information
    p1 = add_paragraph(doc, "1. Tên tổ chức, cá nhân; địa chỉ; điện thoại, fax:", bold=True, space_after=3)
    p1_desc = add_paragraph(doc, "   - Tên tổ chức: ", bold=True, space_after=2)
    add_run(p1_desc, "CÔNG TY TNHH THƯƠNG MẠI XUẤT NHẬP KHẨU EUREKA")
    p1_addr = add_paragraph(doc, "   - Địa chỉ: ", bold=True, space_after=2)
    add_run(p1_addr, "Số 3, hẻm 56, ngõ An Sơn, phố Đại La, phường Trương Định, quận Hai Bà Trưng, thành phố Hà Nội, Việt Nam")
    p1_contact = add_paragraph(doc, "   - Điện thoại: ", bold=True, space_after=8)
    add_run(p1_contact, "024.3999.8888       Fax: 024.3999.8889       Email: contact@eurekalogistics.vn")

    # 2. Product name & details
    add_paragraph(doc, "2. Tên sản phẩm, hàng hoá:", bold=True, space_after=3)
    p2_desc = add_paragraph(doc, "   - Tên hàng hóa: ", bold=True, space_after=2)
    add_run(p2_desc, "Máy sấy tóc dùng trong gia dụng")
    p2_models = add_paragraph(doc, "   - Kiểu, loại/Model: ", bold=True, space_after=2)
    add_run(p2_models, "Model A1, A2, A3 (Chi tiết theo Danh mục sản phẩm kèm theo số: 05/ĐKCBHQ-ERK)")
    p2_brand = add_paragraph(doc, "   - Nhãn hiệu: ", bold=True, space_after=2)
    add_run(p2_brand, "A")
    p2_specs = add_paragraph(doc, "   - Thông số kỹ thuật chính: ", bold=True, space_after=8)
    add_run(p2_specs, "Điện áp: 220V~50Hz; Công suất: 1600W (A1), 1800W (A2), 2000W (A3); Lớp bảo vệ chống điện giật: Lớp II")

    # 3. Assessment location and date
    add_paragraph(doc, "3. Địa điểm, ngày đánh giá:", bold=True, space_after=3)
    p3_desc = add_paragraph(doc, "   - Địa điểm thực hiện đánh giá: ", bold=True, space_after=2)
    add_run(p3_desc, "Kho bảo quản Công ty TNHH Thương mại Xuất nhập khẩu Eureka, Hà Nội")
    p3_date = add_paragraph(doc, "   - Ngày đánh giá: ", bold=True, space_after=8)
    add_run(p3_date, "02 tháng 07 năm 2026")

    # 4. Applicable Standards
    add_paragraph(doc, "4. Số hiệu tiêu chuẩn/quy chuẩn kỹ thuật áp dụng:", bold=True, space_after=3)
    p4_qcvn4 = add_paragraph(doc, "   - ", space_after=2)
    add_run(p4_qcvn4, "QCVN 04:2009/BKHCN và Sửa đổi 1:2016 QCVN 04:2009/BKHCN", bold=True)
    add_run(p4_qcvn4, " - Quy chuẩn kỹ thuật quốc gia về an toàn đối với thiết bị điện và điện tử.")
    p4_qcvn9 = add_paragraph(doc, "   - ", space_after=8)
    add_run(p4_qcvn9, "QCVN 09:2012/BKHCN và Sửa đổi 1:2018 QCVN 09:2012/BKHCN", bold=True)
    add_run(p4_qcvn9, " - Quy chuẩn kỹ thuật quốc gia về tương thích điện từ đối với thiết bị điện và điện tử gia dụng và các mục đích tương tự.")

    # 5. Laboratory details
    add_paragraph(doc, "5. Tên tổ chức thử nghiệm sản phẩm, hàng hoá:", bold=True, space_after=3)
    p5_desc = add_paragraph(doc, "   - Tên phòng thử nghiệm: ", bold=True, space_after=2)
    add_run(p5_desc, "Trung tâm Kỹ thuật Tiêu chuẩn Đo lường Chất lượng 1 (QUATEST 1)")
    p5_vilas = add_paragraph(doc, "   - Mã số công nhận/chỉ định: ", bold=True, space_after=8)
    add_run(p5_vilas, "VILAS 023 (Quyết định chỉ định số 567/QĐ-TĐC của Tổng cục Tiêu chuẩn Đo lường Chất lượng)")

    # 6. Evaluation of Test results
    add_paragraph(doc, "6. Đánh giá về kết quả thử nghiệm theo tiêu chuẩn/quy chuẩn kỹ thuật áp dụng:", bold=True, space_after=3)
    p6_desc = add_paragraph(doc, space_after=8)
    p6_desc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p6_desc, "   Căn cứ Phiếu kết quả thử nghiệm số ")
    add_run(p6_desc, "260702/QT1-1", bold=True)
    add_run(p6_desc, " ngày ")
    add_run(p6_desc, "28/06/2026", bold=True)
    add_run(p6_desc, " của Trung tâm Kỹ thuật Tiêu chuẩn Đo lường Chất lượng 1 (QUATEST 1) cấp cho mẫu Máy sấy tóc dùng trong gia dụng, nhãn hiệu A, các model A1, A2, A3.\n")
    add_run(p6_desc, "   Đánh giá: ")
    add_run(p6_desc, "Tất cả các chỉ tiêu thử nghiệm bao gồm: bảo vệ chống chạm vào các bộ phận mang điện, công suất và dòng điện đầu vào, phát nóng, dòng rò và độ bền điện ở nhiệt độ hoạt động, khả năng chống ẩm, độ bền cơ học, dây dẫn nguồn và dây dẫn ngoài (theo QCVN 04:2009) cũng như giới hạn nhiễu điện từ (theo QCVN 09:2012) đều ")
    add_run(p6_desc, "ĐẠT YÊU CẦU", bold=True)
    add_run(p6_desc, " quy định.")

    # 7. Other contents
    add_paragraph(doc, "7. Các nội dung khác (nếu có):", bold=True, space_after=3)
    p7_desc = add_paragraph(doc, space_after=8)
    p7_desc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p7_desc, "   - Biện pháp kiểm soát chất lượng: Duy trì kiểm soát hồ sơ kỹ thuật, hồ sơ xuất nhập khẩu, giám sát tính nguyên vẹn của lô hàng trong quá trình vận chuyển và bốc dỡ. Đảm bảo dán nhãn hợp quy CR và nhãn phụ tiếng Việt đầy đủ trước khi xuất bán ra thị trường.\n")
    add_run(p7_desc, "   - Tần suất đánh giá giám sát định kỳ: 12 tháng/lần theo quy định của hệ thống quản lý ISO 9001:2015 của nhà sản xuất.")

    # 8. Conclusion
    add_paragraph(doc, "8. Kết luận:", bold=True, space_after=3)
    add_paragraph(doc, "   [x]  Sản phẩm, hàng hoá phù hợp tiêu chuẩn/quy chuẩn kỹ thuật.", bold=True, space_after=2)
    add_paragraph(doc, "   [ ]  Sản phẩm, hàng hoá không phù hợp tiêu chuẩn/quy chuẩn kỹ thuật.", bold=True, space_after=12)

    p_commit = add_paragraph(doc, space_after=18)
    p_commit.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p_commit, "CÔNG TY TNHH THƯƠNG MẠI XUẤT NHẬP KHẨU EUREKA cam kết chất lượng sản phẩm, hàng hóa phù hợp với quy chuẩn kỹ thuật, tiêu chuẩn công bố áp dụng và hoàn toàn chịu trách nhiệm trước pháp luật về chất lượng sản phẩm, hàng hóa và kết quả tự đánh giá.")

    add_signature_block(doc, show_reporter=True)
    return doc

# ---------------------------------------------------------------------------
#  MAIN EXECUTION
# ---------------------------------------------------------------------------
def main():
    output_dir = r"d:\AI AGENT THUY\AI agent Kinh doanh\reports"
    os.makedirs(output_dir, exist_ok=True)
    
    files = {
        "BAN_DANG_KY_CONG_BO_HOP_QUY_Kem_Danh_Muc.docx": build_ban_dang_ky,
        "DANH_MUC_SAN_PHAM_KEM_THEO_CBHQ.docx": build_danh_muc_san_pham,
        "BAO_CAO_TU_DANH_GIA_HOP_QUY.docx": build_bao_cao_tu_danh_gia
    }
    
    print("=" * 65)
    print("  AGENT 05 - GENERATING BATCH OF CONFORMITY DOCUMENTS")
    print("=" * 65)
    
    for filename, builder_func in files.items():
        path = os.path.join(output_dir, filename)
        doc = builder_func()
        doc.save(path)
        size_kb = os.path.getsize(path) // 1024
        print(f"  [OK] Generated {filename} ({size_kb} KB)")
        
    print("=" * 65)
    print(f"  All documents saved successfully in: {output_dir}")
    print("=" * 65)

if __name__ == "__main__":
    main()
