# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

"""
Xuất hợp đồng Eureka ra Word (.docx) — ĐÚNG THEO MẪU CHUẨN EUREKA
Chỉ Auto-Fill thông tin khách hàng, KHÔNG thêm/bớt điều khoản.
Agent 04 - Contract & Finance Assistant v3.0
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Thông tin Auto-Fill ─────────────────────────────────────────────────────
KH = {
    "ten_cty":   "CÔNG TY TNHH THƯƠNG MẠI Y TẾ PHƯƠNG MAI",
    "mst":       "0110888973",
    "dia_chi":   "Số 22/34A/162 phố Đông Thiên, Phường Vĩnh Hưng, TP Hà Nội, Việt Nam",
    "dai_dien":  "Bà Nguyễn Thanh Mai",
    "chuc_vu":   "Giám Đốc",
    "dien_thoai":".................................",
    "email":     ".................................",
    "ngay_ky":   "29 tháng 05 năm 2026",
    "so_hdut":   "01/2026/HĐUT-ERK/PMAI",
    "so_pl":     "01/2026/PLHĐ-ERK/PMAI",
    "ma_kh":     "PMAI",
    "dia_giao":  "Số 22/34A/162 phố Đông Thiên, Phường Vĩnh Hưng, TP Hà Nội, Việt Nam",
    "hieu_luc":  "29/05/2026 đến hết ngày 29/05/2027 (01 năm)",
}

EUREKA = {
    "ten_cty":   "CÔNG TY CỔ PHẦN EUREKA LOGISTICS",
    "mst":       "0109999999",
    "dia_chi":   "Tòa nhà Eureka, Cầu Giấy, Hà Nội",
    "dai_dien":  "Ông Trần Đức Thắng",
    "chuc_vu":   "Tổng Giám Đốc",
    "dien_thoai":"024.xxxx.xxxx",
    "email":     "contact@eurekalogistics.vn",
}

OUTPUT_DIR = r"d:\AI AGENT THUY\AI agent Kinh doanh\reports"

# ─── Màu sắc ─────────────────────────────────────────────────────────────────
BLUE_HEX  = "1A3A6B"
GOLD_HEX  = "C89B2A"
LTBL_HEX  = "EEF3FB"
WHITE_HEX = "FFFFFF"
GRAY_HEX  = "CCCCCC"

FONT      = "Times New Roman"


# ─── Tiện ích ─────────────────────────────────────────────────────────────────

def set_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_margins(doc, t=2.0, b=2.0, l=3.0, r=2.0):
    s = doc.sections[0]
    s.top_margin    = Cm(t)
    s.bottom_margin = Cm(b)
    s.left_margin   = Cm(l)
    s.right_margin  = Cm(r)


def run(para, text, bold=False, italic=False, size=12,
        hex_color=None, underline=False):
    r = para.add_run(text)
    r.bold      = bold
    r.italic    = italic
    r.underline = underline
    r.font.name = FONT
    r.font.size = Pt(size)
    if hex_color:
        r.font.color.rgb = RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )
    return r


def para_center(doc, text, bold=False, italic=False, size=12, hex_color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run(p, text, bold=bold, italic=italic, size=size, hex_color=hex_color)
    return p


def hr(doc, color="999999", width=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), str(width))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, text, size=12.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run(p, text, bold=True, size=size, hex_color=BLUE_HEX)
    return p


def body(doc, text, size=12, justify=True, indent_cm=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.alignment    = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    )
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run(p, text, size=size)
    return p


def sig_table(doc, ten_a, chuc_a, ten_b, chuc_b):
    """Bảng ký tên 2 cột theo mẫu chuẩn Eureka."""
    doc.add_paragraph()
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows_data = [
        ("ĐẠI DIỆN BÊN A",                  "ĐẠI DIỆN BÊN B"),
        ("(Ký, ghi rõ họ tên và đóng dấu)", "(Ký, ghi rõ họ tên và đóng dấu)"),
        ("",                                 ""),
        (ten_a + "\n" + chuc_a,              ten_b + "\n" + chuc_b),
    ]
    bolds = [True, False, False, True]
    for i, (l, r_) in enumerate(rows_data):
        for j, (cell, txt) in enumerate([(tbl.rows[i].cells[0], l),
                                          (tbl.rows[i].cells[1], r_)]):
            set_bg(cell, LTBL_HEX if i == 0 else WHITE_HEX)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(p, txt, bold=bolds[i], size=11.5,
                hex_color=BLUE_HEX if i in (0, 3) else None)
    return tbl


# ═══════════════════════════════════════════════════════════════════════════════
#  1. HỢP ĐỒNG ỦY THÁC NHẬP KHẨU & GOM CONTAINER LCL  (Mẫu chuẩn Eureka)
# ═══════════════════════════════════════════════════════════════════════════════

def build_hd_uy_thac():
    doc = Document()
    set_margins(doc)

    # ── Phần đầu quốc gia ──
    para_center(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
                bold=True, size=13, hex_color=BLUE_HEX)
    para_center(doc, "Độc lập - Tự do - Hạnh phúc",
                bold=True, size=12, hex_color=BLUE_HEX)
    hr(doc, GOLD_HEX, 12)

    # ── Tên hợp đồng ──
    doc.add_paragraph()
    para_center(doc, "HỢP ĐỒNG ỦY THÁC NHẬP KHẨU VÀ GOM CONTAINER LCL",
                bold=True, size=14, hex_color=BLUE_HEX)
    para_center(doc, f"Số: {KH['so_hdut']}",
                bold=True, italic=True, size=12, hex_color=GOLD_HEX)
    doc.add_paragraph()

    # ── Căn cứ ──
    for cc in [
        "Căn cứ Bộ luật Dân sự số 91/2015/QH13 ngày 24/11/2015 và các văn bản hướng dẫn thi hành;",
        "Căn cứ Luật Thương mại số 36/2005/QH11 ngày 14/06/2005 và các văn bản hướng dẫn thi hành;",
        "Căn cứ nhu cầu và khả năng thực tế của hai Bên.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run(p, "- ", bold=True, size=12)
        run(p, cc, italic=True, size=12)

    doc.add_paragraph()
    p_ngay = doc.add_paragraph()
    p_ngay.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run(p_ngay, f"Hôm nay, ngày {KH['ngay_ky']}, tại văn phòng Công ty Cổ phần Eureka Logistics, chúng tôi gồm có:", size=12)

    hr(doc, BLUE_HEX)

    # ── BÊN A (Ủy thác = KH) ──
    heading(doc, "BÊN A: BÊN ỦY THÁC (KHÁCH HÀNG)", size=12.5)
    party_lines_a = [
        f"Tên công ty: {KH['ten_cty']}",
        f"Mã số thuế: {KH['mst']}",
        f"Địa chỉ: {KH['dia_chi']}",
        f"Đại diện: {KH['dai_dien']} - Chức vụ: {KH['chuc_vu']}",
        f"Điện thoại: {KH['dien_thoai']} - Email: {KH['email']}",
    ]
    for line in party_lines_a:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.5)
        run(p, "•  ", bold=True, size=12, hex_color=BLUE_HEX)
        # Split bold key from value
        if ": " in line:
            key, _, val = line.partition(": ")
            run(p, key + ": ", bold=True, size=12)
            run(p, val, size=12)
        else:
            run(p, line, size=12)

    hr(doc, GRAY_HEX, 4)

    # ── BÊN B (Nhận ủy thác = Eureka) ──
    heading(doc, "BÊN B: BÊN NHẬN ỦY THÁC", size=12.5)
    party_lines_b = [
        f"Tên công ty: {EUREKA['ten_cty']}",
        f"Mã số thuế: {EUREKA['mst']}",
        f"Địa chỉ: {EUREKA['dia_chi']}",
        f"Đại diện: {EUREKA['dai_dien']} - Chức vụ: {EUREKA['chuc_vu']}",
        f"Điện thoại: {EUREKA['dien_thoai']} - Email: {EUREKA['email']}",
    ]
    for line in party_lines_b:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.5)
        run(p, "•  ", bold=True, size=12, hex_color=BLUE_HEX)
        if ": " in line:
            key, _, val = line.partition(": ")
            run(p, key + ": ", bold=True, size=12)
            run(p, val, size=12)
        else:
            run(p, line, size=12)

    doc.add_paragraph()
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run(p_intro, "Hai Bên thống nhất ký kết Hợp đồng Ủy thác Nhập khẩu với các điều khoản chi tiết sau:", size=12)

    hr(doc, BLUE_HEX)

    # ── ĐIỀU 1 ──
    heading(doc, "ĐIỀU 1: NỘI DUNG ỦY THÁC")
    clauses_d1 = [
        "1. Bên A ủy thác cho Bên B thực hiện các thủ tục mua hàng, thanh toán quốc tế, làm thủ tục hải quan, gom hàng container lẻ (LCL) và vận chuyển lô hàng từ Trung Quốc/Hàn Quốc về Việt Nam.",
        "2. Bên B đứng tên trên tờ khai hải quan nhập khẩu tại Việt Nam, thực hiện thông quan và giao hàng tận nơi cho Bên A tại địa chỉ yêu cầu.",
        "3. Chi tiết lô hàng được quy định cụ thể tại Phụ lục Hợp đồng đính kèm - là phần không thể tách rời của Hợp đồng này.",
    ]
    for c in clauses_d1:
        body(doc, c)

    # ── ĐIỀU 2 ──
    heading(doc, "ĐIỀU 2: NGUYÊN TẮC THỰC HIỆN HỢP ĐỒNG 2 LỚP")
    clauses_d2 = [
        "1. Để đảm bảo tính hợp pháp của dòng tiền và chứng từ hải quan, Bên B sẽ ký kết Hợp đồng Thương mại Quốc tế với Nhà cung cấp nước ngoài để làm căn cứ chuyển tiền thanh toán quốc tế qua ngân hàng và xin giấy chứng nhận xuất xứ (C/O) hưởng thuế ưu đãi.",
        "2. Bên B cam kết xuất hóa đơn GTGT trả lại hàng nhập khẩu bàn giao cho Bên A và hóa đơn dịch vụ logistics đúng quy định của pháp luật Việt Nam.",
    ]
    for c in clauses_d2:
        body(doc, c)

    # ── ĐIỀU 3 ──
    heading(doc, "ĐIỀU 3: QUYỀN VÀ NGHĨA VỤ CỦA BÊN A (BÊN ỦY THÁC)")
    clauses_d3 = [
        "1. Cung cấp đầy đủ thông tin kỹ thuật, nhãn mác, nguồn gốc xuất xứ của hàng hóa và chịu trách nhiệm hoàn toàn về tính hợp pháp của sản phẩm nhập khẩu tại Việt Nam.",
        "2. Chuyển tiền tạm ứng đầy đủ và đúng hạn cho Bên B theo quy định tại Phụ lục Hợp đồng để Bên B thanh toán tiền hàng và tiền thuế.",
        "3. Tiếp nhận hàng hóa đúng thời gian và địa điểm thỏa thuận, kiểm tra số lượng và chất lượng hàng hóa khi bàn giao.",
    ]
    for c in clauses_d3:
        body(doc, c)

    # ── ĐIỀU 4 ──
    heading(doc, "ĐIỀU 4: QUYỀN VÀ NGHĨA VỤ CỦA BÊN B (BÊN NHẬN ỦY THÁC)")
    clauses_d4 = [
        "1. Đại diện Bên A đàm phán, ký kết hợp đồng ngoại thương, thực hiện chuyển khoản thanh toán quốc tế đúng hạn cho nhà cung cấp.",
        "2. Tổ chức đóng gói, gom cont LCL, vận chuyển và làm thủ tục thông quan hải quan chính ngạch tại Việt Nam.",
        "3. Xuất hóa đơn bàn giao hàng hóa và dịch vụ logistics cho Bên A theo đúng quy định tài chính.",
        "4. Chịu trách nhiệm bảo quản hàng hóa nguyên đai nguyên kiện trong suốt quá trình vận chuyển.",
    ]
    for c in clauses_d4:
        body(doc, c)

    # ── ĐIỀU 5 ──
    heading(doc, "ĐIỀU 5: PHƯƠNG THỨC GIAO NHẬN VÀ KIỂM SOÁT")
    body(doc, f"1. Địa điểm giao hàng: {KH['dia_giao']}.")
    body(doc, "2. Khi nhận hàng, Bên A tiến hành đồng kiểm số lượng thực tế với nhân viên giao nhận của Bên B. Mọi khiếu nại về số lượng/móp méo bên ngoài phải được lập biên bản ngay tại thời điểm giao hàng.")

    # ── ĐIỀU 6 ──
    heading(doc, "ĐIỀU 6: THANH TOÁN VÀ QUYẾT TOÁN CÔNG NỢ")
    body(doc, "1. Tiến độ thanh toán được chia làm 2 đợt chi tiết tại Phụ lục Hợp đồng.")
    body(doc, "2. Phương thức thanh toán: Chuyển khoản vào tài khoản ngân hàng của Bên B.")
    body(doc, "3. Áp dụng quy tắc kiểm soát nợ quá hạn: Nếu Bên A chậm thanh toán Đợt 2 quá hạn quy định, Bên A phải chịu phí phạt trả chậm là 0.05%/ngày trên số tiền chậm nộp. Bên B có quyền áp dụng các biện pháp bảo vệ dòng tiền (giữ hàng bảo đảm) nếu quá hạn thanh toán ≥ 15 ngày.")

    # ── ĐIỀU 7 ──
    heading(doc, "ĐIỀU 7: ĐIỀU KHOẢN CHUNG")
    body(doc, f"1. Hợp đồng có hiệu lực từ ngày {KH['hieu_luc']}.")
    body(doc, "2. Hợp đồng được lập thành 04 (bốn) bản có giá trị pháp lý như nhau, mỗi bên giữ 02 (hai) bản để thực hiện.")

    hr(doc, BLUE_HEX)

    # ── Ký tên ──
    doc.add_paragraph()
    p_loc = doc.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p_loc, f"Hà Nội, ngày {KH['ngay_ky']}", italic=True, bold=True, size=12,
        hex_color=BLUE_HEX)

    sig_table(
        doc,
        ten_a=KH['dai_dien'], chuc_a=KH['chuc_vu'],
        ten_b=EUREKA['dai_dien'], chuc_b=EUREKA['chuc_vu'],
    )

    hr(doc, GOLD_HEX, 8)
    p_f = doc.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p_f,
        f"Mã tài liệu: HĐUT-{KH['so_hdut']} | Ngày xuất: {KH['ngay_ky']} | Agent 04 – Eureka Logistics v3.0",
        italic=True, size=9, hex_color="888888")

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
#  2. PHỤ LỤC HỢP ĐỒNG ỦY THÁC  (Mẫu chuẩn Eureka)
# ═══════════════════════════════════════════════════════════════════════════════

def build_phu_luc():
    doc = Document()
    set_margins(doc)

    para_center(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
                bold=True, size=13, hex_color=BLUE_HEX)
    para_center(doc, "Độc lập - Tự do - Hạnh phúc",
                bold=True, size=12, hex_color=BLUE_HEX)
    hr(doc, GOLD_HEX, 12)

    doc.add_paragraph()
    para_center(doc, "PHỤ LỤC HỢP ĐỒNG ỦY THÁC NHẬP KHẨU LCL",
                bold=True, size=14, hex_color=BLUE_HEX)
    para_center(doc, f"Số: {KH['so_pl']}",
                bold=True, italic=True, size=12, hex_color=GOLD_HEX)
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p_ref,
        f"(Đính kèm Hợp đồng Ủy thác số: {KH['so_hdut']})",
        italic=True, size=11, hex_color="555577")
    doc.add_paragraph()

    # ── Thông tin lô hàng ──
    pl_lines = [
        "Tên hàng hóa: ................................................................ (Việt + Trung)",
        "Số lượng: .................................................................",
        "Quy cách đóng gói: Hàng lẻ đóng ghép container LCL / Hàng nguyên cont FCL",
        "Tuyến đường: .................................................................",
    ]
    for line in pl_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.5)
        run(p, "•  ", bold=True, size=12, hex_color=BLUE_HEX)
        if ": " in line:
            key, _, val = line.partition(": ")
            run(p, key + ": ", bold=True, size=12)
            run(p, val, size=12)
        else:
            run(p, line, size=12)

    hr(doc, BLUE_HEX)

    # ── Bảng chi phí & thuế ──
    heading(doc, "BẢNG TỔNG HỢP CHI PHÍ & THUẾ (ĐỊNH LƯỢNG CHI TIẾT)")

    tbl_headers = ["STT", "Khoản mục chi phí", "Công thức & Cách tính",
                   "Giá trị chưa thuế (VND)", "Thuế suất", "Tiền thuế GTGT (VND)", "Tổng cộng có thuế (VND)"]
    tbl_widths  = [1.0, 4.5, 3.5, 2.2, 1.5, 2.2, 2.2]

    tbl_rows = [
        ["I",      "Phần hàng hóa & Thuế nhập khẩu (Ủy thác hộ)", "", "", "", "", ""],
        ["1",      "Trị giá hàng khai báo", "Trị giá hóa đơn mua hàng quốc tế", "[Trị giá]", "0%", "0", "[Trị giá]"],
        ["2",      "Thuế Nhập khẩu ước tính", "Trị giá hàng × Thuế suất NK (%)", "[Tiền thuế NK]", "0%", "0", "[Tiền thuế NK]"],
        ["3",      "Thuế VAT Hàng nhập khẩu", "(Trị giá hàng + Thuế NK) × 10%", "-", "10%", "[Tiền VAT NK]", "[Tiền VAT NK]"],
        ["",       "Cộng nhóm I (Tạm tính)", "(Ủy thác hộ)", "[Cộng I chưa VAT]", "", "[Tiền VAT NK]", "[Tổng I có VAT]"],
        ["II",     "Phần dịch vụ Logistics & Phí ủy thác", "", "", "", "", ""],
        ["4",      "Cước vận chuyển LCL / FCL", "Đơn giá cước × Thể tích/Trọng lượng", "[Tiền cước]", "10%", "[VAT cước]", "[Tổng cước có VAT]"],
        ["5",      "Phí dịch vụ ủy thác nhập khẩu", "Trị giá hàng × 2% (Phí dịch vụ)", "[Phí dịch vụ]", "10%", "[VAT dịch vụ]", "[Tổng dịch vụ có VAT]"],
        ["",       "Cộng nhóm II (Phí dịch vụ)", "(Doanh thu Eureka)", "[Cộng II chưa VAT]", "", "[Tổng VAT dịch vụ]", "[Tổng II có VAT]"],
        ["TỔNG",  "TỔNG GIÁ TRỊ TOÀN LÔ HÀNG", "Cộng Nhóm I + Nhóm II", "[Tổng chưa VAT]", "", "[Tổng VAT]", "[TỔNG CỘNG LÔ HÀNG]"],
    ]
    bold_rows  = {0, 4, 5, 8, 9}
    group_rows = {0, 5}
    total_rows = {4, 8, 9}

    table = doc.add_table(rows=1 + len(tbl_rows), cols=len(tbl_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for j, (h, w) in enumerate(zip(tbl_headers, tbl_widths)):
        cell = table.rows[0].cells[j]
        cell.width = Cm(w)
        set_bg(cell, BLUE_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, h, bold=True, size=9.5, hex_color="FFFFFF")

    for i, row_data in enumerate(tbl_rows):
        row = table.rows[i + 1]
        is_group = i in group_rows
        is_total = i in total_rows
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.width = Cm(tbl_widths[j])
            bg = LTBL_HEX if is_group else ("D4E4FF" if is_total else WHITE_HEX)
            set_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j not in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            run(p, val, bold=(i in bold_rows), size=9.5,
                hex_color=BLUE_HEX if (is_group or is_total) else None)

    doc.add_paragraph()
    hr(doc, BLUE_HEX)

    # ── Tiến độ thanh toán ──
    heading(doc, "TIẾN ĐỘ THANH TOÁN (LỊCH TRÌNH DÒNG TIỀN BẢO VỆ DOANH NGHIỆP)")

    # Đợt 1
    p_d1 = doc.add_paragraph()
    run(p_d1, "Đợt 1: Tạm ứng trước khi thông quan & mua hàng",
        bold=True, size=12, hex_color="CC0000")

    dot1_lines = [
        ("Số tiền thanh toán:", "[Tiền hàng + Thuế NK + Thuế VAT NK + Cước vận chuyển] VND"),
        ("Chi tiết bao gồm:", "Trị giá tiền hàng (Để Eureka thanh toán quốc tế cho nhà cung cấp); Thuế NK tạm tính và Thuế VAT hàng NK (Để Eureka nộp thuế thông quan); Cước vận chuyển quốc tế tạm ứng trước."),
        ("Thời hạn thanh toán:", "Trong vòng 02 ngày làm việc kể từ ngày ký Phụ lục này (Trước khi hàng xếp lên phương tiện vận chuyển)."),
    ]
    for key, val in dot1_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        run(p, key + " ", bold=True, size=12)
        run(p, val, size=12)

    doc.add_paragraph()

    # Đợt 2
    p_d2 = doc.add_paragraph()
    run(p_d2, "Đợt 2: Tất toán dịch vụ khi bàn giao hàng và Hóa đơn tài chính GTGT đầu vào",
        bold=True, size=12, hex_color=BLUE_HEX)

    dot2_lines = [
        ("Số tiền thanh toán:", "[Phí dịch vụ ủy thác + Thuế VAT cước + Thuế VAT phí dịch vụ] VND"),
        ("Thời hạn thanh toán:", "Ngay khi hàng hóa về tới kho đích của Eureka tại Việt Nam và Bên B phát hành thông báo giao hàng kèm theo bản scan Hóa đơn GTGT dịch vụ đầu vào."),
    ]
    for key, val in dot2_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        run(p, key + " ", bold=True, size=12)
        run(p, val, size=12)

    hr(doc, BLUE_HEX)

    # ── Nguyên tắc xuất hóa đơn ──
    heading(doc, "NGUYÊN TẮC XUẤT HÓA ĐƠN ĐẦU VÀO ĐỂ HỢP THỨC HÓA DOANH NGHIỆP")
    body(doc, "Để Khách hàng hạch toán thuế GTGT và chi phí hợp lệ 100%, Eureka Logistics sẽ phát hành 02 Hóa đơn tài chính như sau:")

    hd_data = [
        ("HÓA ĐƠN 1: BÀN GIAO HÀNG HÓA NHẬP KHẨU ỦY THÁC", [
            f"Tên hàng hóa ghi trên hóa đơn: \"Tên hàng hóa (Trả lại hàng hóa nhập khẩu ủy thác theo HĐ số {KH['so_hdut']})\"",
            "Đơn giá bàn giao: = Giá mua thực tế + Thuế NK thực tế.",
            "Thuế suất GTGT: 10%.",
            "Tiền thuế GTGT: [Tiền thuế GTGT hàng NK thực tế] (Trùng khớp 100% với số thuế GTGT đã nộp tại Hải quan Việt Nam).",
            "Tổng giá trị thanh toán hóa đơn: [Tiền hàng + Thuế NK + Thuế VAT NK] (Đã cấn trừ vào khoản tạm ứng Đợt 1).",
        ]),
        ("HÓA ĐƠN 2: DỊCH VỤ LOGISTICS & PHÍ ỦY THÁC", [
            "Nội dung ghi trên hóa đơn: Phí dịch vụ ủy thác nhập khẩu (2%) và Cước vận chuyển quốc tế.",
            "Thuế suất GTGT: 10%.",
            "Tổng giá trị thanh toán hóa đơn: [Tiền cước + Phí dịch vụ + Thuế VAT dịch vụ] (Đã cấn trừ phần cước tạm ứng ở Đợt 1, khách thanh toán nốt phần còn lại ở Đợt 2).",
        ]),
    ]

    for hd_title, hd_items in hd_data:
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(6)
        p_h.paragraph_format.space_after  = Pt(2)
        p_h.paragraph_format.left_indent  = Cm(0.3)
        run(p_h, hd_title, bold=True, size=12, hex_color=BLUE_HEX)
        for item in hd_items:
            p_i = doc.add_paragraph()
            p_i.paragraph_format.left_indent   = Cm(1.0)
            p_i.paragraph_format.space_before  = Pt(1)
            p_i.paragraph_format.space_after   = Pt(2)
            run(p_i, "• ", bold=True, size=11.5, hex_color=GOLD_HEX)
            if ": " in item:
                k, _, v = item.partition(": ")
                run(p_i, k + ": ", bold=True, size=11.5)
                run(p_i, v, size=11.5)
            else:
                run(p_i, item, size=11.5)

    hr(doc, BLUE_HEX)

    # ── Ký tên ──
    doc.add_paragraph()
    p_loc = doc.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p_loc, "Hà Nội, ngày ......... tháng ......... năm 2026",
        italic=True, bold=True, size=12, hex_color=BLUE_HEX)

    sig_table(
        doc,
        ten_a=KH['dai_dien'], chuc_a=KH['chuc_vu'],
        ten_b=EUREKA['dai_dien'], chuc_b=EUREKA['chuc_vu'],
    )

    hr(doc, GOLD_HEX, 8)
    p_f = doc.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p_f,
        f"Mã tài liệu: {KH['so_pl']} | Ngày xuất: {KH['ngay_ky']} | Agent 04 – Eureka Logistics v3.0",
        italic=True, size=9, hex_color="888888")

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════════

print("=" * 62)
print("  AGENT 04 — XUAT FILE WORD HOP DONG EUREKA (MAU CHUAN)")
print("=" * 62)

files = {
    "HOP_DONG_UY_THAC_PhuongMai_2026.docx": build_hd_uy_thac,
    "PHU_LUC_HOP_DONG_PhuongMai_2026.docx": build_phu_luc,
}

for fname, builder in files.items():
    path = os.path.join(OUTPUT_DIR, fname)
    doc  = builder()
    doc.save(path)
    size_kb = os.path.getsize(path) // 1024
    print(f"  OK  {fname}  ({size_kb} KB)")

print("=" * 62)
print(f"  Thu muc: {OUTPUT_DIR}")
print("  Hoan thanh! Ca 2 file Word da duoc xuat thanh cong.")
print("=" * 62)
