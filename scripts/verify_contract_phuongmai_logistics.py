# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
from docx import Document

REPORTS_DIR = r"D:\AI AGENT THUY\AI agent Kinh doanh\reports"
file_path = os.path.join(REPORTS_DIR, "HOP_DONG_GIAO_NHAN_PhuongMai_2026.docx")

def verify():
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return
        
    doc = Document(file_path)
    print("=" * 60)
    print("VERIFYING GENERATED FORWARDING CONTRACT FOR PHUONG MAI")
    print("=" * 60)
    
    # Contract number
    print(f"Contract Number paragraph: {doc.paragraphs[5].text.strip()}")
    
    # Sign date paragraph
    print(f"Sign Date paragraph: {doc.paragraphs[12].text.strip()}")
    
    # Customer Info
    print("\nCustomer Info paragraphs (Bên A):")
    for idx in range(14, 20):
        print(f"  [P{idx:02d}] {doc.paragraphs[idx].text.strip()}")
        
    # Validity
    print(f"\nValidity paragraph: {doc.paragraphs[84].text.strip()}")
    print(f"Liquidation paragraph: {doc.paragraphs[85].text.strip()}")
    print("=" * 60)

if __name__ == "__main__":
    verify()
