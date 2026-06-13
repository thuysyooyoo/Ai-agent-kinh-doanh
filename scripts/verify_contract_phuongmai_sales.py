# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
from docx import Document

REPORTS_DIR = r"D:\AI AGENT THUY\AI agent Kinh doanh\reports"
file_path = os.path.join(REPORTS_DIR, "HOP_DONG_MUA_BAN_PhuongMai_2026.docx")

def verify():
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return
        
    doc = Document(file_path)
    print("=" * 60)
    print("VERIFYING GENERATED SALES CONTRACT FOR PHUONG MAI")
    print("=" * 60)
    
    # Contract number
    print(f"Contract Number paragraph: {doc.paragraphs[4].text.strip()}")
    
    # Sign date paragraph
    print(f"Sign Date paragraph: {doc.paragraphs[11].text.strip()}")
    
    # Customer Info
    print("\nCustomer Info paragraphs (Bên B):")
    for idx in range(21, 28):
        print(f"  [P{idx:02d}] {doc.paragraphs[idx].text.strip()}")
        
    # Validity
    print(f"\nValidity paragraph: {doc.paragraphs[79].text.strip()}")
    print("=" * 60)

if __name__ == "__main__":
    verify()
