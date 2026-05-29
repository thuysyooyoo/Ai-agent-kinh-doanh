import os
import sys
import pandas as pd
from docx import Document
from pypdf import PdfReader

# Set standard output encoding to UTF-8
sys.stdout.reconfigure(encoding='utf-8')

base_dir = r"D:\AI AGENT THUY\AI agent Kinh doanh\knowledge\customs_rules"

def convert_6_rules():
    src = os.path.join(base_dir, "6 quy tắc.xlsx")
    dest = os.path.join(base_dir, "6_quy_tac.md")
    if not os.path.exists(src):
        print("6 quy tắc.xlsx not found.")
        return
    try:
        df = pd.read_excel(src)
        # Convert to markdown manually without depending on tabulate library
        md_content = "# SÁU QUY TẮC TỔNG QUÁT GIẢI THÍCH HS\n\n"
        headers = [str(c) for c in df.columns]
        md_content += "| " + " | ".join(headers) + " |\n"
        md_content += "| " + " | ".join(["---"] * len(headers)) + " |\n"
        for idx, row in df.iterrows():
            vals = [str(val).replace('\n', '<br>') for val in row.values]
            md_content += "| " + " | ".join(vals) + " |\n"
        with open(dest, "w", encoding="utf-8") as f:
            f.write(md_content)
        print("Successfully converted 6 quy tắc.xlsx to 6_quy_tac.md")
    except Exception as e:
        print(f"Error converting 6 quy tắc: {e}")

def convert_announcements():
    src = os.path.join(base_dir, "Thông báo phân loại.xlsx")
    dest = os.path.join(base_dir, "thong_bao_phan_loai.txt")
    if not os.path.exists(src):
        print("Thông báo phân loại.xlsx not found.")
        return
    try:
        df = pd.read_excel(src)
        # Save as Tab-Separated Values (TSV)
        df.to_csv(dest, sep="\t", index=False, encoding="utf-8")
        print("Successfully converted Thông báo phân loại.xlsx to thong_bao_phan_loai.txt")
    except Exception as e:
        print(f"Error converting announcements: {e}")

def convert_tariff_excel():
    src = os.path.join(base_dir, "Danh mục hàng hóa Xuất nhập khẩu Việt Nam.xlsx")
    dest = os.path.join(base_dir, "danh_muc_xnk.txt")
    if not os.path.exists(src):
        print("Danh mục hàng hóa Xuất nhập khẩu Việt Nam.xlsx not found.")
        return
    try:
        df = pd.read_excel(src, skiprows=1)
        # Save as Tab-Separated Values (TSV)
        df.to_csv(dest, sep="\t", index=False, encoding="utf-8")
        print("Successfully converted Danh mục hàng hóa Xuất nhập khẩu Việt Nam.xlsx to danh_muc_xnk.txt")
    except Exception as e:
        print(f"Error converting tariff excel: {e}")

def convert_docx_commentary():
    src = os.path.join(base_dir, "Bản dịch CHÚ GIẢI CHI TIẾT DANH MỤC HS2022 EN2022.docx")
    dest = os.path.join(base_dir, "chu_giai_chi_tiet_EN2022.txt")
    if not os.path.exists(src):
        print("Bản dịch CHÚ GIẢI CHI TIẾT DANH MỤC HS2022 EN2022.docx not found.")
        return
    try:
        doc = Document(src)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        # Include tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append(" | ".join(row_text))
                
        with open(dest, "w", encoding="utf-8") as f:
            f.write("\n".join(full_text))
        print("Successfully converted Bản dịch CHÚ GIẢI CHI TIẾT DANH MỤC HS2022 EN2022.docx to chu_giai_chi_tiet_EN2022.txt")
    except Exception as e:
        print(f"Error converting docx commentary: {e}")

def convert_pdf_sen():
    src_files = ["Chú giai Sen 2022 bổ sung.pdf", "chu gia sen.pdf"]
    dest = os.path.join(base_dir, "chu_giai_sen_2022.txt")
    
    full_text = []
    for sf in src_files:
        path = os.path.join(base_dir, sf)
        if not os.path.exists(path):
            print(f"{sf} not found.")
            continue
        try:
            print(f"Reading PDF: {sf}...")
            reader = PdfReader(path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    full_text.append(f"--- Page {i+1} of {sf} ---")
                    full_text.append(text)
        except Exception as e:
            print(f"Error reading PDF {sf}: {e}")
            
    if full_text:
        with open(dest, "w", encoding="utf-8") as f:
            f.write("\n".join(full_text))
        print("Successfully extracted PDF text to chu_giai_sen_2022.txt")

if __name__ == "__main__":
    print("Starting conversion of all official files to txt/md...")
    convert_6_rules()
    convert_announcements()
    convert_tariff_excel()
    convert_docx_commentary()
    convert_pdf_sen()
    print("All conversions completed!")
