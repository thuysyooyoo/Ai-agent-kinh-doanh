# -*- coding: utf-8 -*-
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

# Margins: Top: 2cm, Bottom: 2cm, Left: 3cm, Right: 2cm
def set_margins(doc):
    s = doc.sections[0]
    s.top_margin    = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin   = Cm(3.0)
    s.right_margin  = Cm(2.0)

def set_cell_margins(cell, top=0, bottom=0, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_element(name):
    return OxmlElement(name)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    new_borders = create_element('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = create_element(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        new_borders.append(border)
    tblPr.append(new_borders)

FONT_NAME = "Times New Roman"

def add_run(para, text, bold=False, italic=False, size=12, color_rgb=None, underline=False):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    r.font.name = FONT_NAME
    r.font.size = Pt(size)
    if color_rgb:
        r.font.color.rgb = color_rgb
    return r

def add_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False, size=12, space_before=0, space_after=6, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size)
    return p

def main():
    doc = Document()
    set_margins(doc)
    
    # 1. Header block: Importer info & National Emblem (Using a borderless 2-column table)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_table)
    
    # Set widths for columns: left col 6.5cm, right col 9.5cm
    header_table.columns[0].width = Cm(6.5)
    header_table.columns[1].width = Cm(9.5)
    
    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]
    
    # Left Column: Company Name & Document No.
    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.paragraph_format.space_after = Pt(2)
    p_left.paragraph_format.line_spacing = 1.15
    add_run(p_left, "CÔNG TY TNHH THƯƠNG MẠI\nXUẤT NHẬP KHẨU EUREKA", bold=True, size=10.5)
    
    p_no = left_cell.add_paragraph()
    p_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_no.paragraph_format.space_after = Pt(0)
    add_run(p_no, "Số: 05/ĐKCBHQ-ERK", size=11, bold=False)
    
    # Right Column: National Motto
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.paragraph_format.space_after = Pt(2)
    add_run(p_right, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, size=11)
    
    p_motto = right_cell.add_paragraph()
    p_motto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_motto.paragraph_format.space_after = Pt(4)
    add_run(p_motto, "Độc lập - Tự do - Hạnh phúc", bold=True, size=11.5)
    
    # Draw horizontal line under right column motto
    p_line = right_cell.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(0)
    add_run(p_line, "_______________", bold=True, size=11)

    # Spacing
    add_paragraph(doc, "", space_after=18)

    # 2. Document Title
    add_paragraph(doc, "BẢN ĐĂNG KÝ CÔNG BỐ HỢP QUY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=12)

    # 3. Main Info Block
    p_info_intro = add_paragraph(doc, "Tên tổ chức, cá nhân: ", bold=True, space_after=4)
    add_run(p_info_intro, "CÔNG TY TNHH THƯƠNG MẠI XUẤT NHẬP KHẨU EUREKA", bold=True)
    
    p_addr = add_paragraph(doc, "Địa chỉ: ", bold=True, space_after=4)
    add_run(p_addr, "Số 3, hẻm 56, ngõ An Sơn, phố Đại La, phường Trương Định, quận Hai Bà Trưng, thành phố Hà Nội, Việt Nam")
    
    p_contact = add_paragraph(doc, "Điện thoại: ", bold=True, space_after=12)
    add_run(p_contact, "024.3999.8888      ")
    add_run(p_contact, "Fax: ", bold=True)
    add_run(p_contact, "024.3999.8889      ")
    add_run(p_contact, "E-mail: ", bold=True)
    add_run(p_contact, "contact@eurekalogistics.vn")

    # 4. Announcement Section
    add_paragraph(doc, "ĐĂNG KÝ CÔNG BỐ:", bold=True, size=12.5, space_after=6)
    
    p_decl = add_paragraph(doc, "Sản phẩm, hàng hóa, quá trình, dịch vụ, môi trường hoặc các đối tượng khác trong hoạt động kinh tế - xã hội:", italic=True, space_after=6)
    
    p_name = add_paragraph(doc, "- Tên gọi sản phẩm: ", bold=True, space_after=4)
    add_run(p_name, "Máy sấy tóc dùng trong gia dụng")
    
    p_model = add_paragraph(doc, "- Kiểu, loại/Model: ", bold=True, space_after=4)
    add_run(p_model, "Máy sấy tóc A")
    
    p_brand = add_paragraph(doc, "- Nhãn hiệu/Hiệu: ", bold=True, space_after=4)
    add_run(p_brand, "A")
    
    p_specs = add_paragraph(doc, "- Đặc trưng kỹ thuật: ", bold=True, space_after=8)
    add_run(p_specs, "Điện áp: 220V ~ 50Hz; Công suất định mức: 1800W; Lớp bảo vệ chống điện giật: Lớp II")

    # 5. Conformity details
    p_conform = add_paragraph(doc, "Phù hợp với quy chuẩn kỹ thuật:", bold=True, space_after=4)
    p_conform_detail = add_paragraph(doc, "- Số hiệu, ký hiệu, tên gọi: ", bold=True, space_after=8)
    add_run(p_conform_detail, "QCVN 04:2009/BKHCN và Sửa đổi 1:2016 QCVN 04:2009/BKHCN - Quy chuẩn kỹ thuật quốc gia về an toàn đối với thiết bị điện và điện tử")

    # 6. Supplementary Info
    add_paragraph(doc, "Thông tin bổ sung:", bold=True, space_after=4)
    
    p_basis = add_paragraph(doc, "- Căn cứ công bố hợp quy: ", bold=True, space_after=4)
    add_run(p_basis, "Kết quả chứng nhận hợp quy của Tổ chức chứng nhận được chỉ định.")
    
    p_method = add_paragraph(doc, "- Phương thức đánh giá sự phù hợp: ", bold=True, space_after=8)
    add_run(p_method, "Phương thức 5 (đánh giá thử nghiệm mẫu điển hình và giám sát hệ thống quản lý chất lượng) quy định tại Thông tư số 14/2026/TT-BKHCN.")

    # 7. Type of evaluation
    add_paragraph(doc, "Loại hình đánh giá:", bold=True, space_after=4)
    
    p_eval1 = add_paragraph(doc, "   [x]  Tổ chức chứng nhận đánh giá:", bold=True, space_after=2)
    p_eval1_detail = add_paragraph(doc, "         + Tên tổ chức chứng nhận: ", bold=True, space_after=2)
    add_run(p_eval1_detail, "Trung tâm Chứng nhận Phù hợp (QUACERT)")
    p_eval1_cert = add_paragraph(doc, "         + Số giấy chứng nhận: ", bold=True, space_after=2)
    add_run(p_eval1_cert, "26.0702/GCNHQ-QUACERT")
    p_eval1_date = add_paragraph(doc, "         + Ngày cấp: ", bold=True, space_after=6)
    add_run(p_eval1_date, "30 tháng 06 năm 2026")
    
    p_eval2 = add_paragraph(doc, "   [ ]  Tự đánh giá: ", bold=True, space_after=12)
    add_run(p_eval2, "Ngày lãnh đạo tổ chức, cá nhân ký xác nhận Báo cáo tự đánh giá: ........................")

    # 8. Commitment
    p_commit = add_paragraph(doc, space_after=18)
    p_commit.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p_commit, "CÔNG TY TNHH THƯƠNG MẠI XUẤT NHẬP KHẨU EUREKA ", bold=True)
    add_run(p_commit, "cam kết và chịu trách nhiệm về tính phù hợp của sản phẩm, hàng hóa, quá trình, dịch vụ, môi trường và các đối tượng khác trong hoạt động kinh tế - xã hội do mình sản xuất, kinh doanh, bảo quản, vận chuyển, sử dụng, khai thác.")

    # 9. Signatures Block
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(sig_table)
    sig_table.columns[0].width = Cm(7.0)
    sig_table.columns[1].width = Cm(9.0)
    
    s_left = sig_table.rows[0].cells[0]
    s_right = sig_table.rows[0].cells[1]
    
    # Right Column signature
    p_date = s_right.paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_after = Pt(2)
    add_run(p_date, "Hà Nội, ngày 02 tháng 07 năm 2026", italic=True)
    
    p_rep = s_right.add_paragraph()
    p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rep.paragraph_format.space_after = Pt(2)
    add_run(p_rep, "ĐẠI DIỆN TỔ CHỨC, CÁ NHÂN", bold=True)
    
    p_role = s_right.add_paragraph()
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role.paragraph_format.space_after = Pt(40)  # Large space for signature
    add_run(p_role, "(Ký tên, chức vụ, đóng dấu)", italic=True, size=11)
    
    p_name_sign = s_right.add_paragraph()
    p_name_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name_sign.paragraph_format.space_after = Pt(0)
    add_run(p_name_sign, "Trần Đức Thắng", bold=True)

    # Save to reports directory
    output_dir = r"d:\AI AGENT THUY\AI agent Kinh doanh\reports"
    os.makedirs(output_dir, exist_ok=True)
    filename = "BAN_DANG_KY_CONG_BO_HOP_QUY_May_Say_Toc_A.docx"
    output_path = os.path.join(output_dir, filename)
    
    doc.save(output_path)
    print(f"SUCCESS: Generated conformity document at: {output_path}")

if __name__ == "__main__":
    main()
