# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
import os
from docx import Document

TEMPLATE_DIR = r"D:\AI AGENT THUY\AI agent Kinh doanh\templates\biểu mẫu giấy tờ\2. CHỨNG TỪ KHÁCH HÀNG ỦY THÁC"
OUTPUT_FILE = r"D:\AI AGENT THUY\AI agent Kinh doanh\reports\inspect_templates_full.txt"

files = [
    "1 HỢP ĐỒNG THƯƠNG MẠI MUA BÁN HÀNG HÓA.docx",
    "3. HỢP ĐỒNG ỦY THÁC GOM CONT.docx"
]

with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
    for fname in files:
        path = os.path.join(TEMPLATE_DIR, fname)
        f.write(f"\n{'='*80}\n")
        f.write(f"FILE: {fname}\n")
        f.write(f"{'='*80}\n")
        
        if not os.path.exists(path):
            f.write(f"Error: File not found at {path}\n")
            continue
            
        doc = Document(path)
        f.write(f"--- PARAGRAPHS ({len(doc.paragraphs)}) ---\n")
        for i, para in enumerate(doc.paragraphs):
            t = para.text
            f.write(f"[P{i:03d}] {repr(t)}\n")
            
        f.write(f"\n--- TABLES ({len(doc.tables)}) ---\n")
        for ti, tbl in enumerate(doc.tables):
            f.write(f"\n--- TABLE {ti} (Rows: {len(tbl.rows)}, Cols: {len(tbl.rows[0].cells) if tbl.rows else 0}) ---\n")
            for ri, row in enumerate(tbl.rows):
                # We want to print cell content by cell index to map it properly
                row_cells = []
                for ci, cell in enumerate(row.cells):
                    ct = cell.text.strip()
                    row_cells.append(f"[C{ci}] {repr(ct)}")
                f.write(f"  [Row {ri}] {' | '.join(row_cells)}\n")

print(f"Inspection complete. Written to {OUTPUT_FILE}")
