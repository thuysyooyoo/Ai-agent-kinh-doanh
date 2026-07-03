# -*- coding: utf-8 -*-
import os
import sys
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

# Page Margins: Top: 2cm, Bottom: 2cm, Left: 3cm, Right: 2cm
def set_margins(doc):
    s = doc.sections[0]
    s.top_margin    = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin   = Cm(3.0)
    s.right_margin  = Cm(2.0)

def create_element(name):
    return OxmlElement(name)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    new_borders = create_element('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = create_element(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        new_borders.append(border)
    tblPr.append(new_borders)

FONT_NAME = "Times New Roman"

def add_run(para, text, bold=False, italic=False, size=12, color_rgb=None, underline=False):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    r.font.name = FONT_NAME
    r.font.size = Pt(size)
    if color_rgb:
        r.font.color.rgb = color_rgb
    return r

def add_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False, size=12, space_before=0, space_after=6, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size)
    return p

def add_national_header(doc, company_name, doc_number):
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_table)
    
    header_table.columns[0].width = Cm(7.0)
    header_table.columns[1].width = Cm(9.0)
    
    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]
    
    # Left Column: Company Name & Document No.
    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.paragraph_format.space_after = Pt(2)
    p_left.paragraph_format.line_spacing = 1.15
    add_run(p_left, company_name, bold=True, size=10.0)
    
    p_no = left_cell.add_paragraph()
    p_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_no.paragraph_format.space_after = Pt(0)
    add_run(p_no, f"Số: {doc_number}", size=10.5, bold=False)
    
    # Right Column: National Motto
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.paragraph_format.space_after = Pt(2)
    add_run(p_right, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, size=11)
    
    p_motto = right_cell.add_paragraph()
    p_motto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_motto.paragraph_format.space_after = Pt(4)
    add_run(p_motto, "Độc lập - Tự do - Hạnh phúc", bold=True, size=11.5)
    
    # Line under right column motto
    p_line = right_cell.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(0)
    add_run(p_line, "_______________", bold=True, size=11)

def add_signature_block(doc, config, show_reporter=False):
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(sig_table)
    sig_table.columns[0].width = Cm(7.5)
    sig_table.columns[1].width = Cm(8.5)
    
    s_left = sig_table.rows[0].cells[0]
    s_right = sig_table.rows[0].cells[1]
    
    if show_reporter:
        p_rep_title = s_left.paragraphs[0]
        p_rep_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rep_title.paragraph_format.space_after = Pt(2)
        add_run(p_rep_title, "NGƯỜI ĐÁNH GIÁ", bold=True)
        
        p_rep_sign = s_left.add_paragraph()
        p_rep_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rep_sign.paragraph_format.space_after = Pt(40)
        add_run(p_rep_sign, "(Ký và ghi rõ họ tên)", italic=True, size=11)
        
        p_rep_name = s_left.add_paragraph()
        p_rep_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p_rep_name, "Nguyễn Văn Hùng", bold=True)

    p_date = s_right.paragraphs[0] if not show_reporter else s_right.paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_after = Pt(2)
    add_run(p_date, f"Hà Nội, ngày {config['declaration_date']}", italic=True)
    
    p_rep = s_right.add_paragraph()
    p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rep.paragraph_format.space_after = Pt(2)
    add_run(p_rep, "XÁC NHẬN CỦA LÃNH ĐẠO" if show_reporter else "ĐẠI DIỆN TỔ CHỨC, CÁ NHÂN", bold=True)
    
    p_role = s_right.add_paragraph()
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role.paragraph_format.space_after = Pt(40)
    add_run(p_role, "(Ký tên, chức vụ, đóng dấu)", italic=True, size=11)
    
    p_name_sign = s_right.add_paragraph()
    p_name_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name_sign.paragraph_format.space_after = Pt(0)
    add_run(p_name_sign, config['importer']['representative'], bold=True)

# ---------------------------------------------------------------------------
#  GENERATE DANG KY CBHQ
# ---------------------------------------------------------------------------
def generate_dang_ky_cbhq(config, qcvn_key, qcvn_name, qcvn_desc, cert_info, output_path):
    doc = Document()
    set_margins(doc)
    
    doc_no = f"{config[qcvn_key]['doc_no_prefix']}"
    add_national_header(doc, config['importer']['name_short'], doc_no)
    
    add_paragraph(doc, "", space_after=18)
    add_paragraph(doc, "BẢN ĐĂNG KÝ CÔNG BỐ HỢP QUY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=12)

    p_info_intro = add_paragraph(doc, "Tên tổ chức, cá nhân: ", bold=True, space_after=4)
    add_run(p_info_intro, config['importer']['name'], bold=True)
    
    p_addr = add_paragraph(doc, "Địa chỉ: ", bold=True, space_after=4)
    add_run(p_addr, config['importer']['address'])
    
    p_contact = add_paragraph(doc, "Điện thoại: ", bold=True, space_after=12)
    add_run(p_contact, f"{config['importer']['phone']}      ")
    add_run(p_contact, "Fax: ", bold=True)
    add_run(p_contact, f"{config['importer']['fax']}      ")
    add_run(p_contact, "E-mail: ", bold=True)
    add_run(p_contact, config['importer']['email'])

    add_paragraph(doc, "CÔNG BỐ:", bold=True, size=12.5, space_after=6)
    add_paragraph(doc, "Sản phẩm, hàng hóa, quá trình, dịch vụ, môi trường hoặc các đối tượng khác trong hoạt động kinh tế - xã hội:", italic=True, space_after=6)
    
    p_name = add_paragraph(doc, "- Tên gọi sản phẩm: ", bold=True, space_after=4)
    add_run(p_name, config['product']['name_vi'])
    
    p_model = add_paragraph(doc, "- Kiểu, loại/Model: ", bold=True, space_after=4)
    add_run(p_model, "Chi tiết tại Danh mục sản phẩm kèm theo (Phụ lục số 01)", italic=True)
    
    p_brand = add_paragraph(doc, "- Nhãn hiệu/Hiệu: ", bold=True, space_after=4)
    add_run(p_brand, "Chi tiết tại Danh mục sản phẩm kèm theo (Phụ lục số 01)", italic=True)
    
    p_specs = add_paragraph(doc, "- Đặc trưng kỹ thuật: ", bold=True, space_after=8)
    add_run(p_specs, "Chi tiết tại Danh mục sản phẩm kèm theo (Phụ lục số 01)", italic=True)

    p_conform = add_paragraph(doc, "Phù hợp với quy chuẩn kỹ thuật:", bold=True, space_after=4)
    p_conform_detail = add_paragraph(doc, "- Số hiệu, ký hiệu, tên gọi: ", bold=True, space_after=8)
    add_run(p_conform_detail, f"{qcvn_name} - {qcvn_desc}")

    add_paragraph(doc, "Thông tin bổ sung:", bold=True, space_after=4)
    p_basis = add_paragraph(doc, "- Căn cứ công bố hợp quy: ", bold=True, space_after=4)
    add_run(p_basis, cert_info['basis'])
    
    p_method = add_paragraph(doc, "- Phương thức đánh giá sự phù hợp: ", bold=True, space_after=8)
    add_run(p_method, cert_info['method'])

    add_paragraph(doc, "Loại hình đánh giá:", bold=True, space_after=4)
    
    is_cert = cert_info['type'] == 'cert'
    chk_cert = "[x]" if is_cert else "[ ]"
    chk_self = "[x]" if not is_cert else "[ ]"
    
    p_eval1 = add_paragraph(doc, f"   {chk_cert}  Tổ chức chứng nhận đánh giá:", bold=True, space_after=2)
    p_eval1_detail = add_paragraph(doc, "         + Tên tổ chức chứng nhận: ", bold=True, space_after=2)
    add_run(p_eval1_detail, config[qcvn_key].get('cert_body', '........................'))
    p_eval1_cert = add_paragraph(doc, "         + Số giấy chứng nhận: ", bold=True, space_after=2)
    add_run(p_eval1_cert, config[qcvn_key].get('cert_no', '........................'))
    p_eval1_date = add_paragraph(doc, "         + Ngày cấp: ", bold=True, space_after=6)
    add_run(p_eval1_date, config[qcvn_key].get('cert_date', '........................'))
    
    p_eval2 = add_paragraph(doc, f"   {chk_self}  Tự đánh giá: ", bold=True, space_after=12)
    self_date = config['declaration_date'] if not is_cert else "........................"
    add_run(p_eval2, f"Ngày lãnh đạo tổ chức, cá nhân ký xác nhận Báo cáo tự đánh giá: {self_date}")

    p_commit = add_paragraph(doc, space_after=18)
    p_commit.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p_commit, f"{config['importer']['name']} ", bold=True)
    add_run(p_commit, "cam kết và chịu trách nhiệm về tính phù hợp của sản phẩm, hàng hóa, quá trình, dịch vụ, môi trường và các đối tượng khác trong hoạt động kinh tế - xã hội do mình sản xuất, kinh doanh, bảo quản, vận chuyển, sử dụng, khai thác.")

    add_signature_block(doc, config, show_reporter=False)
    doc.save(output_path)

# ---------------------------------------------------------------------------
#  GENERATE BAO CAO TU DANH GIA
# ---------------------------------------------------------------------------
def generate_bao_cao_tu_danh_gia(config, qcvn_key, qcvn_name, qcvn_desc, cert_info, output_path):
    doc = Document()
    set_margins(doc)
    
    doc_no = f"{config[qcvn_key]['doc_no_prefix'].replace('ĐKCBHQ', 'BC')}"
    add_national_header(doc, config['importer']['name_short'], doc_no)
    
    add_paragraph(doc, "", space_after=18)
    add_paragraph(doc, "BÁO CÁO TỰ ĐÁNH GIÁ HỢP QUY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=12)

    # 1. Company Information
    add_paragraph(doc, "1. Tên tổ chức, cá nhân; địa chỉ; điện thoại, fax:", bold=True, space_after=3)
    p1_desc = add_paragraph(doc, "   - Tên tổ chức: ", bold=True, space_after=2)
    add_run(p1_desc, config['importer']['name'])
    p1_addr = add_paragraph(doc, "   - Địa chỉ: ", bold=True, space_after=2)
    add_run(p1_addr, config['importer']['address'])
    p1_contact = add_paragraph(doc, "   - Điện thoại: ", bold=True, space_after=8)
    add_run(p1_contact, f"{config['importer']['phone']}       Fax: {config['importer']['fax']}       Email: {config['importer']['email']}")

    # 2. Product name & details
    add_paragraph(doc, "2. Tên sản phẩm, hàng hoá:", bold=True, space_after=3)
    p2_desc = add_paragraph(doc, "   - Tên hàng hóa: ", bold=True, space_after=2)
    add_run(p2_desc, config['product']['name_vi'])
    p2_models = add_paragraph(doc, "   - Kiểu, loại/Model: ", bold=True, space_after=2)
    add_run(p2_models, f"Chi tiết tại Danh mục sản phẩm kèm theo (Phụ lục số 01)")
    p2_brand = add_paragraph(doc, "   - Nhãn hiệu: ", bold=True, space_after=2)
    add_run(p2_brand, config['product']['brand'])
    p2_specs = add_paragraph(doc, "   - Thông số kỹ thuật chính: ", bold=True, space_after=8)
    add_run(p2_specs, "Chi tiết tại Danh mục sản phẩm kèm theo (Phụ lục số 01)", italic=True)

    # 3. Assessment location and date
    add_paragraph(doc, "3. Địa điểm, ngày đánh giá:", bold=True, space_after=3)
    p3_desc = add_paragraph(doc, "   - Địa điểm thực hiện đánh giá: ", bold=True, space_after=2)
    add_run(p3_desc, f"Kho bảo quản {config['importer']['name']}")
    p3_date = add_paragraph(doc, "   - Ngày đánh giá: ", bold=True, space_after=8)
    add_run(p3_date, config['declaration_date'])

    # 4. Applicable Standards
    add_paragraph(doc, "4. Số hiệu tiêu chuẩn/quy chuẩn kỹ thuật áp dụng:", bold=True, space_after=3)
    p4_qcvn = add_paragraph(doc, "   - ", space_after=8)
    add_run(p4_qcvn, qcvn_name, bold=True)
    add_run(p4_qcvn, f" - {qcvn_desc}")

    # 5. Laboratory details
    add_paragraph(doc, "5. Tên tổ chức thử nghiệm sản phẩm, hàng hoá:", bold=True, space_after=3)
    p5_desc = add_paragraph(doc, "   - Tên phòng thử nghiệm: ", bold=True, space_after=2)
    add_run(p5_desc, config[qcvn_key]['lab_name'])
    p5_vilas = add_paragraph(doc, "   - Mã số công nhận/chỉ định: ", bold=True, space_after=8)
    add_run(p5_vilas, "VILAS 023 (Chỉ định số 567/QĐ-TĐC của Tổng cục Tiêu chuẩn Đo lường Chất lượng)")

    # 6. Evaluation of Test results
    add_paragraph(doc, "6. Đánh giá về kết quả thử nghiệm theo tiêu chuẩn/quy chuẩn kỹ thuật áp dụng:", bold=True, space_after=3)
    p6_desc = add_paragraph(doc, space_after=8)
    p6_desc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if cert_info['type'] == 'cert':
        add_run(p6_desc, "   Căn cứ Giấy chứng nhận hợp quy số ")
        add_run(p6_desc, config[qcvn_key].get('cert_no', '................'), bold=True)
        add_run(p6_desc, " ngày ")
        add_run(p6_desc, config[qcvn_key].get('cert_date', '................'), bold=True)
        add_run(p6_desc, f" do {config[qcvn_key].get('cert_body', '................')} cấp (trên cơ sở Phiếu kết quả thử nghiệm số ")
        add_run(p6_desc, config[qcvn_key]['test_report_no'], bold=True)
        add_run(p6_desc, " ngày ")
        add_run(p6_desc, config[qcvn_key]['test_report_date'], bold=True)
        add_run(p6_desc, f" của {config[qcvn_key]['lab_name']} cấp cho mẫu {config['product']['name_vi']}).\n")
        add_run(p6_desc, f"   Đánh giá: Lô sản phẩm nhập khẩu đã được Tổ chức chứng nhận được chỉ định đánh giá phù hợp chất lượng và cấp Giấy chứng nhận hợp quy hợp lệ, đáp ứng đầy đủ các yêu cầu kỹ thuật quy định tại ")
        add_run(p6_desc, qcvn_name, bold=True)
        add_run(p6_desc, ".")
    else:
        add_run(p6_desc, "   Căn cứ Phiếu kết quả thử nghiệm số ")
        add_run(p6_desc, config[qcvn_key]['test_report_no'], bold=True)
        add_run(p6_desc, " ngày ")
        add_run(p6_desc, config[qcvn_key]['test_report_date'], bold=True)
        add_run(p6_desc, f" của {config[qcvn_key]['lab_name']} cấp cho mẫu {config['product']['name_vi']}.\n")
        add_run(p6_desc, f"   Đánh giá: Tất cả các chỉ tiêu thử nghiệm được yêu cầu quy định tại ")
        add_run(p6_desc, qcvn_name, bold=True)
        add_run(p6_desc, " đều đạt kết quả thử nghiệm ")
        add_run(p6_desc, "ĐẠT YÊU CẦU", bold=True)
        add_run(p6_desc, " và phù hợp với giới hạn quy định.")

    # 7. Other contents
    add_paragraph(doc, "7. Các nội dung khác (nếu có):", bold=True, space_after=3)
    p7_desc = add_paragraph(doc, space_after=8)
    p7_desc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p7_desc, "   - Biện pháp kiểm soát chất lượng: Duy trì kiểm soát hồ sơ kỹ thuật, hồ sơ nhập khẩu hải quan chính ngạch. Thực hiện dán nhãn hợp quy CR và nhãn phụ tiếng Việt đầy đủ trước khi lưu thông.\n")
    add_run(p7_desc, "   - Kế hoạch giám sát: Lấy mẫu thử nghiệm và kiểm soát chất lượng định kỳ hàng năm.")

    # 8. Conclusion
    add_paragraph(doc, "8. Kết luận:", bold=True, space_after=3)
    add_paragraph(doc, "   [x]  Sản phẩm, hàng hoá phù hợp tiêu chuẩn/quy chuẩn kỹ thuật.", bold=True, space_after=2)
    add_paragraph(doc, "   [ ]  Sản phẩm, hàng hoá không phù hợp tiêu chuẩn/quy chuẩn kỹ thuật.", bold=True, space_after=12)

    p_commit = add_paragraph(doc, space_after=18)
    p_commit.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p_commit, f"{config['importer']['name']} cam kết chất lượng sản phẩm, hàng hóa phù hợp với quy chuẩn kỹ thuật, tiêu chuẩn công bố áp dụng và hoàn toàn chịu trách nhiệm trước pháp luật về chất lượng sản phẩm, hàng hóa và kết quả tự đánh giá.")

    add_signature_block(doc, config, show_reporter=True)
    doc.save(output_path)

# ---------------------------------------------------------------------------
#  GENERATE DANH MUC SAN PHAM (No Weight column)
# ---------------------------------------------------------------------------
def generate_danh_muc(config, qcvn_key, qcvn_name, output_path):
    doc = Document()
    set_margins(doc)
    
    doc_no_dk = f"{config[qcvn_key]['doc_no_prefix']}"
    doc_no_bc = f"{doc_no_dk.replace('ĐKCBHQ', 'BC')}"
    
    add_national_header(doc, config['importer']['name_short'], f"Kèm theo số: {doc_no_dk}")
    add_paragraph(doc, "", space_after=12)
    
    # Add 'PHỤ LỤC SỐ 01' at the top right of the catalog document
    p_pl = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)
    add_run(p_pl, "PHỤ LỤC SỐ 01", bold=True, size=11)
    
    add_paragraph(doc, f"DANH MỤC SẢN PHẨM {config['product']['name_vi'].upper()}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_after=2)
    
    p_sub = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_run(p_sub, f"(Kèm theo bản đăng ký công bố hợp quy số: {doc_no_dk} và báo cáo tự đánh giá số: {doc_no_bc})", bold=True, italic=True, size=10.5)

    # 9 columns (removed Weight column)
    tbl_headers = [
        "TT", "Tên sản phẩm, hàng hóa", "Nhãn hiệu", "Mã HS", "Xuất xứ", 
        "Số lượng (cái)", "Đặc tính kỹ thuật", "Nhóm sản phẩm", "Quy chuẩn áp dụng"
    ]
    tbl_widths = [0.8, 4.2, 1.3, 1.5, 1.3, 1.3, 3.2, 1.2, 2.5]
    
    table = doc.add_table(rows=1 + len(config['product']['models']), cols=len(tbl_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header styling
    for col_idx, (text, width) in enumerate(zip(tbl_headers, tbl_widths)):
        cell = table.rows[0].cells[col_idx]
        cell.width = Cm(width)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, text, bold=True, size=9.5)
        
    # Row styling
    for row_idx, model_data in enumerate(config['product']['models']):
        row = table.rows[row_idx + 1]
        
        row_data = [
            str(row_idx + 1),
            f"{config['product']['name_vi']}, Model {model_data['model']}, mới 100%",
            config['product']['brand'],
            config['product']['hs_code'],
            config['product']['origin'],
            model_data['qty'],
            model_data['specs'],
            "Nhóm 2",
            qcvn_name
        ]
        
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = Cm(tbl_widths[col_idx])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.0
            
            if col_idx in [0, 2, 3, 4, 5, 7, 8]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            add_run(p, text, size=9.0)
            
    add_paragraph(doc, "", space_after=18)
    
    # Signature
    sig_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(sig_table)
    sig_table.columns[0].width = Cm(8.0)
    sig_table.columns[1].width = Cm(8.0)
    
    s_right = sig_table.rows[0].cells[1]
    
    p_rep = s_right.paragraphs[0]
    p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rep.paragraph_format.space_after = Pt(2)
    add_run(p_rep, "ĐẠI DIỆN TỔ CHỨC, CÁNHÂN", bold=True, size=11)
    
    p_role = s_right.add_paragraph()
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role.paragraph_format.space_after = Pt(40)
    add_run(p_role, "(Ký tên, chức vụ, đóng dấu)", italic=True, size=10)
    
    p_name = s_right.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_name, config['importer']['representative'], bold=True, size=11)
    
    doc.save(output_path)

# ---------------------------------------------------------------------------
#  MAIN PROGRAM
# ---------------------------------------------------------------------------
def main():
    # Load config file if it exists, otherwise use defaults
    config_path = r"d:\AI AGENT THUY\AI agent Kinh doanh\scratch\conformity_config.json"
    
    if os.path.exists(config_path):
        print(f"Loading custom config from: {config_path}")
        with open(config_path, "r", encoding="utf-8") as f:
            config = json.load(f)
    else:
        # Default fallback config
        print("Using default config values.")
        config = {
            "importer": {
                "name": "CÔNG TY TNHH THƯƠNG MẠI XUẤT NHẬP KHẨU EUREKA",
                "name_short": "CÔNG TY TNHH TM\nXUẤT NHẬP KHẨU EUREKA",
                "address": "Số 3, hẻm 56, ngõ An Sơn, phố Đại La, phường Trương Định, quận Hai Bà Trưng, thành phố Hà Nội, Việt Nam",
                "phone": "024.3999.8888",
                "fax": "024.3999.8889",
                "email": "contact@eurekalogistics.vn",
                "representative": "Trần Đức Thắng",
                "role": "Giám đốc"
            },
            "product": {
                "name_vi": "Máy sấy tóc dùng trong gia dụng",
                "hs_code": "8516.31.00",
                "brand": "A",
                "origin": "Trung Quốc",
                "models": [
                    {"model": "A1", "specs": "Điện áp: 220V~50Hz; Công suất: 1600W", "qty": "500"},
                    {"model": "A2", "specs": "Điện áp: 220V~50Hz; Công suất: 1800W", "qty": "300"},
                    {"model": "A3", "specs": "Điện áp: 220V~50Hz; Công suất: 2000W", "qty": "200"}
                ]
            },
            "qcvn_04": {
                "cert_body": "Trung tâm Chứng nhận Phù hợp (QUACERT)",
                "cert_no": "26.0702/GCNHQ-QUACERT",
                "cert_date": "30 tháng 06 năm 2026",
                "lab_name": "Trung tâm Kỹ thuật Tiêu chuẩn Đo lường Chất lượng 1 (QUATEST 1)",
                "test_report_no": "260702/QT1-SAFE",
                "test_report_date": "28/06/2026",
                "doc_no_prefix": "05/ĐKCBHQ-04/ERK"
            },
            "qcvn_09": {
                "lab_name": "Trung tâm Kỹ thuật Tiêu chuẩn Đo lường Chất lượng 1 (QUATEST 1)",
                "test_report_no": "260702/QT1-EMC",
                "test_report_date": "28/06/2026",
                "doc_no_prefix": "05/ĐKCBHQ-09/ERK"
            },
            "declaration_date": "02 tháng 07 năm 2026"
        }
    
    output_dir = r"d:\AI AGENT THUY\AI agent Kinh doanh\reports"
    os.makedirs(output_dir, exist_ok=True)
    
    # 1. Generate files for QCVN 04
    method_04_text = config['qcvn_04'].get(
        'method_text', 
        "Phương thức 7 (thử nghiệm, đánh giá lô sản phẩm) quy định tại Thông tư số 14/2026/TT-BKHCN."
    )
    basis_04_text = f"Kết quả chứng nhận hợp quy của Tổ chức chứng nhận được chỉ định ({config['qcvn_04'].get('cert_body', 'QUACERT')})."
    
    qcvn_04_info = {
        "basis": basis_04_text,
        "method": method_04_text,
        "type": "cert"
    }
    
    path_dk_04 = os.path.join(output_dir, "BAN_DANG_KY_CONG_BO_HOP_QUY_QCVN_04.docx")
    path_bc_04 = os.path.join(output_dir, "BAO_CAO_TU_DANH_GIA_HOP_QUY_QCVN_04.docx")
    path_dm_04 = os.path.join(output_dir, "DANH_MUC_SAN_PHAM_QCVN_04.docx")
    
    generate_dang_ky_cbhq(config, "qcvn_04", "QCVN 04:2009/BKHCN và Sửa đổi 1:2016 QCVN 04:2009/BKHCN", "Quy chuẩn kỹ thuật quốc gia về an toàn đối với thiết bị điện và điện tử", qcvn_04_info, path_dk_04)
    generate_bao_cao_tu_danh_gia(config, "qcvn_04", "QCVN 04:2009/BKHCN và Sửa đổi 1:2016 QCVN 04:2009/BKHCN", "Quy chuẩn kỹ thuật quốc gia về an toàn đối với thiết bị điện và điện tử", qcvn_04_info, path_bc_04)
    generate_danh_muc(config, "qcvn_04", "QCVN 04:2009/BKHCN", path_dm_04)
    
    # 2. Generate files for QCVN 09 (EMC - Medium Risk - Method 1)
    qcvn_09_info = {
        "basis": "Kết quả tự đánh giá hợp quy của tổ chức, cá nhân.",
        "method": "Phương thức 1 (thử nghiệm mẫu điển hình) quy định tại Thông tư số 14/2026/TT-BKHCN.",
        "type": "self"
    }
    
    path_dk_09 = os.path.join(output_dir, "BAN_DANG_KY_CONG_BO_HOP_QUY_QCVN_09.docx")
    path_bc_09 = os.path.join(output_dir, "BAO_CAO_TU_DANH_GIA_HOP_QUY_QCVN_09.docx")
    path_dm_09 = os.path.join(output_dir, "DANH_MUC_SAN_PHAM_QCVN_09.docx")
    
    generate_dang_ky_cbhq(config, "qcvn_09", "QCVN 09:2012/BKHCN và Sửa đổi 1:2018 QCVN 09:2012/BKHCN", "Quy chuẩn kỹ thuật quốc gia về tương thích điện từ đối với thiết bị điện và điện tử gia dụng và các mục đích tương tự", qcvn_09_info, path_dk_09)
    generate_bao_cao_tu_danh_gia(config, "qcvn_09", "QCVN 09:2012/BKHCN và Sửa đổi 1:2018 QCVN 09:2012/BKHCN", "Quy chuẩn kỹ thuật quốc gia về tương thích điện từ đối với thiết bị điện và điện tử gia dụng và các mục đích tương tự", qcvn_09_info, path_bc_09)
    generate_danh_muc(config, "qcvn_09", "QCVN 09:2012/BKHCN", path_dm_09)
    
    print("FINISHED batch generation of QCVN 04 and QCVN 09 conformity documentation.")
    print("Files successfully generated:")
    print("QCVN 04 (Safety):")
    print(f" - {path_dk_04}")
    print(f" - {path_bc_04}")
    print(f" - {path_dm_04}")
    print("QCVN 09 (EMC):")
    print(f" - {path_dk_09}")
    print(f" - {path_bc_09}")
    print(f" - {path_dm_09}")

if __name__ == "__main__":
    main()
