# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
from docx import Document
from docx.shared import Pt

TEMPLATE_DIR = r"D:\AI AGENT THUY\AI agent Kinh doanh\templates\biểu mẫu giấy tờ\1. CHỨNG TỪ KHÁCH HÀNG TỰ ĐỨNG TÊN"
OUTPUT_DIR   = r"D:\AI AGENT THUY\AI agent Kinh doanh\reports"

# Ensure output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_para_text(para, text):
    """
    Safely reset paragraph text while preserving font style.
    """
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        r = para.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

def fill_hdgn_phuongmai():
    src_path = os.path.join(TEMPLATE_DIR, "EUREKA- Hop dong giao nhan van tai .docx")
    dst_path = os.path.join(OUTPUT_DIR, "HOP_DONG_GIAO_NHAN_PhuongMai_2026.docx")
    
    if not os.path.exists(src_path):
        print(f"Error: Template file does not exist at {src_path}")
        return
        
    doc = Document(src_path)
    
    # [P005] Số hợp đồng
    set_para_text(doc.paragraphs[5], "Hợp đồng số: ERK/PM-01022026")
    
    # [P012] Ngày ký hợp đồng
    set_para_text(doc.paragraphs[12], "Hôm nay, ngày 01 tháng 02 năm 2026 tại Hà Nội, chúng tôi gồm:")
    
    # [P014] Bên A: BÊN GIAO NHẬN (KHÁCH HÀNG)
    set_para_text(doc.paragraphs[14], "Bên A: CÔNG TY TNHH THƯƠNG MẠI Y TẾ PHƯƠNG MAI")
    
    # [P015] Địa chỉ
    set_para_text(doc.paragraphs[15], "Địa chỉ: Số 22/34A/162 phố Đông Thiên, Phường Vĩnh Hưng, TP Hà Nội, Việt Nam")
    
    # [P016] Điện thoại
    set_para_text(doc.paragraphs[16], "Điện thoại: .................................")
    
    # [P017] MST
    set_para_text(doc.paragraphs[17], "MST: 0110888973")
    
    # [P018] Người đại diện
    set_para_text(doc.paragraphs[18], "Người đại diện: Ông Nguyễn Thanh Mai         Chức vụ: Giám đốc")
    
    # [P019] Số tài khoản / Điện thoại phụ
    set_para_text(doc.paragraphs[19], "Tài khoản: .................................")
    
    # [P084] Hiệu lực hợp đồng
    set_para_text(doc.paragraphs[84], "7.3. Hợp đồng này có hiệu lực từ ngày 01 tháng 02 năm 2026 đến ngày 01 tháng 02 năm 2028")
    
    # [P085] Xóa ngày thanh lý hợp đồng lỗi thời hoặc viết lại hợp lý
    set_para_text(doc.paragraphs[85], "Hai bên sẽ họp và lập biên bản thanh lý hợp đồng vận chuyển hàng hóa này khi hết hiệu lực hợp đồng.")
    
    doc.save(dst_path)
    print(f"Successfully generated: {dst_path}")

if __name__ == "__main__":
    fill_hdgn_phuongmai()
