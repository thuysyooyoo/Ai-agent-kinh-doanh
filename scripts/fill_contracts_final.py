# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
from docx import Document
from docx.shared import Pt

# ─── Đường dẫn thư mục ───────────────────────────────────────────────────────
TEMPLATE_DIR = r"D:\AI AGENT THUY\AI agent Kinh doanh\templates\biểu mẫu giấy tờ\2. CHỨNG TỪ KHÁCH HÀNG ỦY THÁC"
OUTPUT_DIR   = r"D:\AI AGENT THUY\AI agent Kinh doanh\reports"

# Tạo thư mục output nếu chưa có
os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_para_text(para, text):
    """
    Đặt lại văn bản cho một paragraph một cách an toàn.
    Giữ nguyên định dạng của run đầu tiên (font, size, bold) nếu có.
    """
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        r = para.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

# ─────────────────────────────────────────────────────────────────────────────
#  1. HỢP ĐỒNG THƯƠNG MẠI MUA BÁN HÀNG HÓA
# ─────────────────────────────────────────────────────────────────────────────
def fill_hdtm():
    src_path = os.path.join(TEMPLATE_DIR, "1 HỢP ĐỒNG THƯƠNG MẠI MUA BÁN HÀNG HÓA.docx")
    dst_path = os.path.join(OUTPUT_DIR, "HOP_DONG_THUONG_MAI_PhuongMai_2026_Final.docx")
    
    if not os.path.exists(src_path):
        print(f"Error: Mẫu HĐTM không tồn tại tại {src_path}")
        return
        
    doc = Document(src_path)
    
    # [P004] Số hợp đồng
    set_para_text(doc.paragraphs[4], "Số: 01/2026/HĐTM-ERK/PMAI")
    # [P005] Xóa dòng ví dụ
    set_para_text(doc.paragraphs[5], "")
    
    # [P011] Ngày ký hợp đồng
    set_para_text(doc.paragraphs[11], "Hôm nay ngày 29 tháng 05 năm 2026, tại Văn phòng CÔNG TY TNHH XUẤT NHẬP KHẨU VÀ THƯƠNG MẠI EUREKA")
    
    # [P021] Bên B: BÊN MUA (giữ nguyên)
    # [P022] Cá nhân: ... -> Công ty: CÔNG TY TNHH THƯƠNG MẠI Y TẾ PHƯƠNG MAI
    set_para_text(doc.paragraphs[22], "Công ty: CÔNG TY TNHH THƯƠNG MẠI Y TẾ PHƯƠNG MAI")
    
    # [P023] Địa chỉ
    set_para_text(doc.paragraphs[23], "Địa chỉ: Số 22/34A/162 phố Đông Thiên, Phường Vĩnh Hưng, TP Hà Nội, Việt Nam")
    
    # [P024] Điện thoại (để trống cho khách tự điền nếu phát sinh)
    set_para_text(doc.paragraphs[24], "Điện thoại: .................................")
    
    # [P025] Số CMND -> Thay thế hoàn toàn bằng Mã số thuế (Doanh nghiệp không dùng CMND/CCCD)
    set_para_text(doc.paragraphs[25], "Mã số thuế: 0110888973")
    
    # [P026] Số tài khoản -> Thay thế hoàn toàn bằng Người đại diện & Chức vụ (Xóa dòng tài khoản theo yêu cầu)
    set_para_text(doc.paragraphs[26], "Người đại diện: Ông Nguyễn Thanh Mai         Chức vụ: Giám Đốc")
    
    # [P027] Email
    set_para_text(doc.paragraphs[27], "Email: .................................")
    
    # [P079] Hiệu lực hợp đồng (1 năm)
    set_para_text(doc.paragraphs[79], "    Hợp đồng này có hiệu lực từ ngày 29 tháng 05 năm 2026 đến ngày 29 tháng 05 năm 2027. Sau thời hạn này nếu không bên nào có khiếu nại nào , hợp đồng sẽ tự động được thanh lý.")
    
    doc.save(dst_path)
    print(f"  [OK] Đã xuất file HĐTM: HOP_DONG_THUONG_MAI_PhuongMai_2026_Final.docx")

# ─────────────────────────────────────────────────────────────────────────────
#  2. HỢP ĐỒNG ỦY THÁC GOM CONTAINER LCL
# ─────────────────────────────────────────────────────────────────────────────
def fill_hdut():
    src_path = os.path.join(TEMPLATE_DIR, "3. HỢP ĐỒNG ỦY THÁC GOM CONT.docx")
    dst_path = os.path.join(OUTPUT_DIR, "HOP_DONG_UY_THAC_GOM_CONT_PhuongMai_2026_Final.docx")
    
    if not os.path.exists(src_path):
        print(f"Error: Mẫu HĐUT không tồn tại tại {src_path}")
        return
        
    doc = Document(src_path)
    
    # [P004] Số hợp đồng
    set_para_text(doc.paragraphs[4], "Số: 01/2026/HĐUT-ERK/PMAI")
    # [P005] Xóa dòng ví dụ
    set_para_text(doc.paragraphs[5], "")
    
    # [P007] Ngày ký hợp đồng
    set_para_text(doc.paragraphs[7], " \t\tHà Nội, ngày 29 tháng 05 năm 2026")
    
    # [P008] BÊN A: BÊN ỦY THÁC
    set_para_text(doc.paragraphs[8], "BÊN A: BÊN ỦY THÁC (KHÁCH HÀNG)")
    
    # [P009] Đại diện là -> Thay thế hoàn toàn bằng Công ty
    set_para_text(doc.paragraphs[9], "Công ty: CÔNG TY TNHH THƯƠNG MẠI Y TẾ PHƯƠNG MAI")
    
    # [P010] Địa chỉ -> Thay thế hoàn toàn bằng Mã số thuế
    set_para_text(doc.paragraphs[10], "Mã số thuế: 0110888973")
    
    # [P011] Điện thoại -> Thay thế hoàn toàn bằng Địa chỉ
    set_para_text(doc.paragraphs[11], "Địa chỉ: Số 22/34A/162 phố Đông Thiên, Phường Vĩnh Hưng, TP Hà Nội, Việt Nam")
    
    # [P012] Số căn cước công dân -> Thay thế hoàn toàn bằng Người đại diện & Chức vụ (Xóa hoàn toàn CCCD)
    set_para_text(doc.paragraphs[12], "Người đại diện: Ông Nguyễn Thanh Mai         Chức vụ: Giám Đốc")
    
    # Chèn thêm 2 dòng Điện thoại và Email để thông tin Bên A đầy đủ và chuyên nghiệp
    p_tel = doc.paragraphs[12].insert_paragraph_before("Điện thoại: .................................")
    p_tel.runs[0].font.name = 'Times New Roman'
    p_tel.runs[0].font.size = Pt(12)
    
    p_email = doc.paragraphs[13].insert_paragraph_before("Email: .................................")
    p_email.runs[0].font.name = 'Times New Roman'
    p_email.runs[0].font.size = Pt(12)
    
    # Cập nhật hiệu lực hợp đồng thông minh bằng chuỗi tìm kiếm
    found_eff = False
    for p in doc.paragraphs:
        if "có hiệu lực từ ngày" in p.text and "đến ngày" in p.text:
            set_para_text(p, "7.3. Hợp đồng này có hiệu lực từ ngày 29 tháng 05 năm 2026 đến ngày 29 tháng 05 năm 2027")
            found_eff = True
            break
            
    if not found_eff:
        # Fallback về index nếu không tìm thấy chuỗi
        set_para_text(doc.paragraphs[94], "7.3. Hợp đồng này có hiệu lực từ ngày 29 tháng 05 năm 2026 đến ngày 29 tháng 05 năm 2027")

    doc.save(dst_path)
    print(f"  [OK] Đã xuất file HĐUT: HOP_DONG_UY_THAC_GOM_CONT_PhuongMai_2026_Final.docx")

# ─────────────────────────────────────────────────────────────────────────────
#  MAIN EXECUTION
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 70)
    print("  AGENT 04 - EXECUTE CONTRACT FILLING FOR PHUONG MAI CO.")
    print("=" * 70)
    
    fill_hdtm()
    fill_hdut()
    
    print("=" * 70)
    print("  Done! 2 main contracts generated in reports/ folder.")
    print("  Appendix contract (Phu_luc_HDUT) is skipped as requested.")
    print("=" * 70)
