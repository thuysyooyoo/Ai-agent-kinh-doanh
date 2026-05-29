# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
from docx import Document

REPORTS_DIR = r"D:\AI AGENT THUY\AI agent Kinh doanh\reports"

hdtm_file = os.path.join(REPORTS_DIR, "HOP_DONG_THUONG_MAI_PhuongMai_2026_Final.docx")
hdut_file = os.path.join(REPORTS_DIR, "HOP_DONG_UY_THAC_GOM_CONT_PhuongMai_2026_Final.docx")

def verify_hdtm():
    print("\n" + "="*50)
    print("XÁC THỰC HỢP ĐỒNG THƯƠNG MẠI MUA BÁN HÀNG HÓA (HĐTM)")
    print("="*50)
    
    if not os.path.exists(hdtm_file):
        print("LỖI: Không tìm thấy file HĐTM!")
        return False
        
    doc = Document(hdtm_file)
    success = True
    
    # 1. Kiểm tra Số HĐTM
    so_hd = doc.paragraphs[4].text
    print(f"[-] Số Hợp đồng: {so_hd}")
    if "01/2026/HĐTM-ERK/PMAI" not in so_hd:
        print(" -> [LỖI] Số hợp đồng không chính xác!")
        success = False
        
    # 2. Kiểm tra thông tin Bên Mua (Bên B)
    print("[-] Thông tin Bên B (Bên Mua):")
    party_b_indices = range(22, 28)
    for idx in party_b_indices:
        t = doc.paragraphs[idx].text.strip()
        print(f"    [P{idx:02d}] {t}")
        
        # Check CMND/CCCD and Bank Account deletion
        if "CMND" in t or "CCCD" in t or "căn cước" in t or "hộ chiếu" in t:
            print(" -> [LỖI] Vẫn còn trường CMND/CCCD!")
            success = False
        if "tài khoản" in t and "Bên B" in t:
            print(" -> [LỖI] Vẫn còn trường Số tài khoản ngân hàng của Bên B!")
            success = False
            
        # Check 'Ông' representative name
        if "Nguyễn Thanh Mai" in t and "Bà" in t:
            print(" -> [LỖI] Dùng sai danh xưng 'Bà'!")
            success = False
            
    # 3. Kiểm tra ngày hiệu lực
    eff_text = doc.paragraphs[79].text
    print(f"[-] Hiệu lực hợp đồng: {eff_text.strip()}")
    if "29 tháng 05 năm 2026" not in eff_text or "29 tháng 05 năm 2027" not in eff_text:
        print(" -> [LỖI] Ngày hiệu lực không chính xác!")
        success = False
        
    return success

def verify_hdut():
    print("\n" + "="*50)
    print("XÁC THỰC HỢP ĐỒNG ỦY THÁC GOM CONT (HĐUT)")
    print("="*50)
    
    if not os.path.exists(hdut_file):
        print("LỖI: Không tìm thấy file HĐUT!")
        return False
        
    doc = Document(hdut_file)
    success = True
    
    # 1. Kiểm tra Số HĐUT
    so_hd = doc.paragraphs[4].text
    print(f"[-] Số Hợp đồng: {so_hd}")
    if "01/2026/HĐUT-ERK/PMAI" not in so_hd:
        print(" -> [LỖI] Số hợp đồng không chính xác!")
        success = False
        
    # 2. Kiểm tra thông tin Bên A (Bên Ủy thác)
    print("[-] Thông tin Bên A (Bên Ủy thác):")
    # Chúng ta đã chèn 2 paragraph mới ở P12, do đó phạm vi Bên A là từ P8 đến P15
    for idx in range(8, 16):
        t = doc.paragraphs[idx].text.strip()
        print(f"    [P{idx:02d}] {t}")
        
        # Check CMND/CCCD and Bank Account deletion
        if "CMND" in t or "CCCD" in t or "căn cước" in t or "hộ chiếu" in t:
            print(" -> [LỖI] Vẫn còn trường CMND/CCCD của Bên A!")
            success = False
            
        # Check 'Ông' representative name
        if "Nguyễn Thanh Mai" in t and "Bà" in t:
            print(" -> [LỖI] Dùng sai danh xưng 'Bà'!")
            success = False
            
    # 3. Kiểm tra ngày hiệu lực
    eff_text = ""
    for p in doc.paragraphs:
        if "có hiệu lực từ ngày" in p.text:
            eff_text = p.text.strip()
            break
    print(f"[-] Hiệu lực hợp đồng: {eff_text}")
    if "29 tháng 05 năm 2026" not in eff_text or "29 tháng 05 năm 2027" not in eff_text:
        print(" -> [LỖI] Ngày hiệu lực không chính xác!")
        success = False
        
    return success

if __name__ == "__main__":
    hdtm_ok = verify_hdtm()
    hdut_ok = verify_hdut()
    
    print("\n" + "="*50)
    if hdtm_ok and hdut_ok:
        print("KẾT QUẢ: TẤT CẢ FILE ĐỀU ĐẠT YÊU CẦU 100%!")
    else:
        print("KẾT QUẢ: CÓ LỖI XẢY RA, CẦN KIỂM TRA LẠI!")
    print("="*50 + "\n")
