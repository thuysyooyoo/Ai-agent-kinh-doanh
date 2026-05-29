# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

"""
Script điền thông tin KH vào MẪU GỐC Eureka (.docx)
- Không thêm/bớt/sửa bất kỳ điều khoản nào
- Chỉ thay thế placeholder bằng thông tin KH thực tế
"""

import os, shutil, re
from docx import Document
from docx.shared import Pt
from copy import deepcopy

# ─── Thư mục ─────────────────────────────────────────────────────────────────
TEMPLATE_DIR = r"D:\AI AGENT THUY\AI agent Kinh doanh\templates\biểu mẫu giấy tờ\2. CHỨNG TỪ KHÁCH HÀNG ỦY THÁC"
OUTPUT_DIR   = r"D:\AI AGENT THUY\AI agent Kinh doanh\reports"

# ─── Thông tin khách hàng Phương Mai ─────────────────────────────────────────
# Mapping: text cũ (mẫu) → text mới (KH thực tế)
# Dùng cho cả HĐTM và HĐUT
REPLACE_MAP = {
    # Ngày tháng
    "ngày .... tháng .... năm 20....":   "ngày 29 tháng 05 năm 2026",
    "ngày … tháng … năm 20…":           "ngày 29 tháng 05 năm 2026",
    "ngày…….tháng…….năm 20……":         "ngày 29 tháng 05 năm 2026",
    "ngày    tháng    năm 20":           "ngày 29 tháng 05 năm 2026",
    "05 tháng  03 năm 2020":            "29 tháng 05 năm 2026",
    "05  tháng 03 năm 2021":            "29 tháng 05 năm 2027",
    "18 tháng 07 năm 2020":             "29 tháng 05 năm 2026",

    # Số hợp đồng
    "Số :……/HĐUT-ERK/……":              "Số: 01/2026/HĐUT-ERK/PMAI",
    "Số: ……/HĐTM-ERK/……":             "Số: 01/2026/HĐTM-ERK/PMAI",
    "số ……………":                       "số 01/2026/HĐUT-ERK/PMAI",

    # Thông tin BÊN A (Khách hàng)
    "[Tên Công ty BÊN A]":              "CÔNG TY TNHH THƯƠNG MẠI Y TẾ PHƯƠNG MAI",
    "[Tên công ty bên A]":              "CÔNG TY TNHH THƯƠNG MẠI Y TẾ PHƯƠNG MAI",
    "Tên Công ty:":                     "Tên Công ty: CÔNG TY TNHH THƯƠNG MẠI Y TẾ PHƯƠNG MAI",
    "[MST Bên A]":                      "0110888973",
    "[Địa chỉ Bên A]":                  "Số 22/34A/162 phố Đông Thiên, Phường Vĩnh Hưng, TP Hà Nội",
    "[Người đại diện Bên A]":           "Bà Nguyễn Thanh Mai",
    "[Chức vụ Bên A]":                  "Giám Đốc",

    # Thời hạn hợp đồng
    "từ ngày  05 tháng  03 năm 2020  đến ngày 05  tháng 03 năm 2021":
        "từ ngày 29 tháng 05 năm 2026 đến ngày 29 tháng 05 năm 2027",

    # Địa chỉ giao hàng
    "[Địa chỉ giao hàng]":              "Số 22/34A/162 phố Đông Thiên, Phường Vĩnh Hưng, TP Hà Nội",
}

# Placeholder dạng regex (khớp linh hoạt hơn)
REGEX_REPLACE = [
    # Bên A – tên công ty (dòng trống hoặc dấu chấm)
    (r"(?i)(Công\s*ty\s*[:：]\s*)[\.\…]{3,}", r"\1CÔNG TY TNHH THƯƠNG MẠI Y TẾ PHƯƠNG MAI"),
    (r"(?i)(Tên\s*doanh\s*nghiệp\s*[:：]\s*)[\.\…]{3,}", r"\1CÔNG TY TNHH THƯƠNG MẠI Y TẾ PHƯƠNG MAI"),
    # MST
    (r"(?i)(Mã\s*số\s*thuế\s*[:：]\s*)[\.\…_]{3,}", r"\g<1>0110888973"),
    # Địa chỉ
    (r"(?i)(Địa\s*chỉ\s*[:：]\s*)[\.\…_]{3,}", r"\g<1>Số 22/34A/162 phố Đông Thiên, Phường Vĩnh Hưng, TP Hà Nội"),
    # Người đại diện
    (r"(?i)(Người\s*đại\s*diện\s*[:：]\s*)[\.\…_]{3,}", r"\g<1>Bà Nguyễn Thanh Mai"),
    (r"(?i)(Đại\s*diện\s*[:：]\s*)[\.\…_]{5,}", r"\g<1>Bà Nguyễn Thanh Mai"),
    # Chức vụ
    (r"(?i)(Chức\s*vụ\s*[:：]\s*)[\.\…_]{3,}", r"\g<1>Giám Đốc"),
    # Điện thoại bên A (để trống dấu chấm)
    (r"(?i)(Điện\s*thoại\s*[:：]\s*)[\.\…_]{4,}(\s*-?\s*Email)", r"\g<1>.................................\g<2>"),
    # Số CMND / tài khoản BÊN A — để trống vì chưa có
    (r"(?i)(Số\s*CMND\s*[:：]\s*)$", r"\g<1>................................"),
    (r"(?i)(Số\s*tài\s*khoản\s*[:：]\s*)$", r"\g<1>................................"),
]


def replace_in_paragraph(para):
    """Thay thế text trong paragraph, giữ nguyên formatting."""
    full_text = para.text

    # --- Exact string replace ---
    new_text = full_text
    for old, new in REPLACE_MAP.items():
        if old in new_text:
            new_text = new_text.replace(old, new)

    # --- Regex replace ---
    for pattern, repl in REGEX_REPLACE:
        new_text = re.sub(pattern, repl, new_text)

    if new_text == full_text:
        return  # Không có gì thay đổi

    # Ghi lại vào run đầu tiên, xóa các run còn lại
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""


def replace_in_table(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_in_paragraph(para)


def fill_template(src_path: str, dst_path: str):
    doc = Document(src_path)

    # Xử lý tất cả paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para)

    # Xử lý tất cả bảng
    for table in doc.tables:
        replace_in_table(table)

    # Xử lý header/footer nếu có
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_paragraph(para)
        for para in section.footer.paragraphs:
            replace_in_paragraph(para)

    doc.save(dst_path)


# ─── MAIN ────────────────────────────────────────────────────────────────────

files_to_fill = [
    ("1 HỢP ĐỒNG THƯƠNG MẠI MUA BÁN HÀNG HÓA.docx",
     "HOP_DONG_THUONG_MAI_PhuongMai_2026.docx"),

    ("3. HỢP ĐỒNG ỦY THÁC GOM CONT.docx",
     "HOP_DONG_UY_THAC_GOM_CONT_PhuongMai_2026.docx"),

    ("4. Phụ lục HĐUT.docx",
     "PHU_LUC_HDUT_PhuongMai_2026.docx"),
]

print("=" * 62)
print("  AGENT 04 - DIEN THONG TIN KH VAO MAU GOC EUREKA")
print("=" * 62)

for src_name, dst_name in files_to_fill:
    src = os.path.join(TEMPLATE_DIR, src_name)
    dst = os.path.join(OUTPUT_DIR, dst_name)
    fill_template(src, dst)
    size_kb = os.path.getsize(dst) // 1024
    print(f"  OK  {dst_name}  ({size_kb} KB)")

print("=" * 62)
print(f"  Thu muc: {OUTPUT_DIR}")
print("  Hoan thanh! 3 file da duoc dien thong tin va luu vao reports/")
print("=" * 62)
