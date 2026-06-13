# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
from docx import Document
from docx.shared import Pt

TEMPLATE_DIR = r"D:\AI AGENT THUY\AI agent Kinh doanh\templates\biểu mẫu giấy tờ\2. CHỨNG TỪ KHÁCH HÀNG ỦY THÁC"
OUTPUT_DIR   = r"D:\AI AGENT THUY\AI agent Kinh doanh\reports"

# Ensure output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_para_text(para, text):
    """
    Safely reset the text of a paragraph while preserving the font style.
    """
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        r = para.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

def fill_hdtm_tram():
    src_path = os.path.join(TEMPLATE_DIR, "1 HỢP ĐỒNG THƯƠNG MẠI MUA BÁN HÀNG HÓA.docx")
    dst_path = os.path.join(OUTPUT_DIR, "HOP_DONG_MUA_BAN_TranNgocTram_2026.docx")
    
    if not os.path.exists(src_path):
        print(f"Error: Template file does not exist at {src_path}")
        return
        
    doc = Document(src_path)
    
    # [P004] Số hợp đồng
    set_para_text(doc.paragraphs[4], "Số: ERK/TRAM-03062026")
    # [P005] Xóa dòng ví dụ
    set_para_text(doc.paragraphs[5], "")
    
    # [P011] Ngày ký hợp đồng
    set_para_text(doc.paragraphs[11], "Hôm nay ngày 03 tháng 06 năm 2026, tại Văn phòng CÔNG TY TNHH XUẤT NHẬP KHẨU VÀ THƯƠNG MẠI EUREKA")
    
    # [P022] Bên B: BÊN MUA
    set_para_text(doc.paragraphs[22], "Hộ kinh doanh: HỘ KINH DOANH- TRẦN NGỌC TRÂM")
    
    # [P023] Địa chỉ
    set_para_text(doc.paragraphs[23], "Địa chỉ: Số nhà 70, đường La Tinh 2, thôn La Tinh, Xã An Khánh, Thành phố Hà Nội, Việt Nam")
    
    # [P024] Điện thoại
    set_para_text(doc.paragraphs[24], "Điện thoại: .................................")
    
    # [P025] CCCD & MST
    set_para_text(doc.paragraphs[25], "Số CCCD: 037196006289     Mã số thuế: 037196006289")
    
    # [P026] Người đại diện & Chức vụ
    set_para_text(doc.paragraphs[26], "Người đại diện: Bà Trần Ngọc Trâm         Chức vụ: Đại diện Hộ kinh doanh")
    
    # [P027] Email
    set_para_text(doc.paragraphs[27], "Email: .................................")
    
    # [P079] Hiệu lực hợp đồng
    set_para_text(doc.paragraphs[79], "    Hợp đồng này có hiệu lực từ ngày 03 tháng 06 năm 2026 đến ngày 03 tháng 06 năm 2030. Sau thời hạn này nếu không bên nào có khiếu nại nào , hợp đồng sẽ tự động được thanh lý.")
    
    doc.save(dst_path)
    print(f"Successfully generated: {dst_path}")

if __name__ == "__main__":
    fill_hdtm_tram()
