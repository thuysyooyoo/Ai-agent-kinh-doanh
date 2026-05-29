# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import os

TEMPLATE_DIR = r"D:\AI AGENT THUY\AI agent Kinh doanh\templates\biểu mẫu giấy tờ\2. CHỨNG TỪ KHÁCH HÀNG ỦY THÁC"

files = {
    "HDTM": "1 HỢP ĐỒNG THƯƠNG MẠI MUA BÁN HÀNG HÓA.docx",
    "HDUT": "3. HỢP ĐỒNG ỦY THÁC GOM CONT.docx",
    "PLUT": "4. Phụ lục HĐUT.docx",
}

for key, fname in files.items():
    path = os.path.join(TEMPLATE_DIR, fname)
    print(f"\n{'='*70}")
    print(f"FILE: {fname}")
    print('='*70)
    doc = Document(path)
    for i, para in enumerate(doc.paragraphs):
        t = para.text
        if t.strip():
            print(f"[P{i:03d}] {repr(t)}")
    for ti, tbl in enumerate(doc.tables):
        print(f"\n--- TABLE {ti+1} ---")
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                ct = cell.text.strip()
                if ct:
                    print(f"  [R{ri}C{ci}] {repr(ct)}")
